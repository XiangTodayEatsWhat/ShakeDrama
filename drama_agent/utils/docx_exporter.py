"""
DOCX文档导出器
"""
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from typing import Optional

from ..models import Bible, Episode


class DocxExporter:
    """DOCX文档导出器"""
    
    def __init__(self):
        self.doc = None
    
    def export(self, bible: Bible, output_path: str) -> str:
        """
        导出Bible为DOCX文档
        
        Args:
            bible: 世界观圣经
            output_path: 输出路径
        
        Returns:
            输出文件路径
        """
        self.doc = Document()
        
        # 设置文档样式
        self._setup_styles()
        
        # 添加标题
        self._add_title(bible)
        
        # 添加基本信息
        self._add_basic_info(bible)
        
        # 添加角色介绍
        self._add_characters(bible)
        
        # 添加剧本正文
        self._add_episodes(bible)
        
        # 保存文档
        self.doc.save(output_path)
        return output_path
    
    def _setup_styles(self):
        """设置文档样式"""
        # 设置默认字体
        style = self.doc.styles['Normal']
        font = style.font
        font.name = '宋体'
        font.size = Pt(12)
    
    def _add_title(self, bible: Bible):
        """添加文档标题"""
        title = self.doc.add_heading(bible.title, level=0)
        title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        
        # 添加副标题
        subtitle = self.doc.add_paragraph()
        subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = subtitle.add_run(f"{', '.join(bible.genre)} | {bible.target_audience}")
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(128, 128, 128)
        
        self.doc.add_paragraph()  # 空行
    
    def _add_basic_info(self, bible: Bible):
        """添加基本信息"""
        self.doc.add_heading('剧情梗概', level=1)
        
        p = self.doc.add_paragraph()
        p.add_run(bible.synopsis)
        p.paragraph_format.first_line_indent = Inches(0.5)
        
        self.doc.add_paragraph()
        
        # 添加主题
        if bible.theme:
            p = self.doc.add_paragraph()
            p.add_run('主题：').bold = True
            p.add_run(bible.theme)
        
        # 添加集数信息
        p = self.doc.add_paragraph()
        p.add_run('总集数：').bold = True
        p.add_run(f"{bible.total_episodes}集")
        
        self.doc.add_page_break()
    
    def _add_characters(self, bible: Bible):
        """添加角色介绍"""
        if not bible.characters:
            return
        
        self.doc.add_heading('主要角色', level=1)
        
        for name, char in bible.characters.items():
            # 角色名
            self.doc.add_heading(name, level=2)
            
            # 身份
            p = self.doc.add_paragraph()
            p.add_run('身份：').bold = True
            p.add_run(char.identity)
            
            # 性格
            if char.personality:
                p = self.doc.add_paragraph()
                p.add_run('性格：').bold = True
                p.add_run(char.personality)
            
            # 背景
            if char.background:
                p = self.doc.add_paragraph()
                p.add_run('背景：').bold = True
                p.add_run(char.background)
            
            # 技能
            if char.skills:
                p = self.doc.add_paragraph()
                p.add_run('技能：').bold = True
                p.add_run(', '.join(char.skills))
            
            self.doc.add_paragraph()
        
        self.doc.add_page_break()
    
    def _add_episodes(self, bible: Bible):
        """添加剧本正文"""
        if not bible.episodes:
            return
        
        self.doc.add_heading('剧本正文', level=1)
        
        for episode in sorted(bible.episodes, key=lambda x: x.number):
            # 集数标题
            self.doc.add_heading(f'第{episode.number}集', level=2)
            
            # 本集梗概
            if episode.synopsis:
                p = self.doc.add_paragraph()
                run = p.add_run(episode.synopsis)
                run.italic = True
                run.font.color.rgb = RGBColor(100, 100, 100)
                self.doc.add_paragraph()
            
            # 剧本内容
            script_lines = episode.full_script.split('\n')
            for line in script_lines:
                line = line.strip()
                if not line:
                    self.doc.add_paragraph()
                    continue
                
                # 场景头（包含数字-数字格式，支持markdown的##前缀）
                if self._is_scene_header(line):
                    # 去掉markdown的##和#前缀（如果有）
                    clean_line = line.lstrip('#').strip()
                    p = self.doc.add_paragraph()
                    run = p.add_run(clean_line)
                    run.bold = True
                    run.font.size = Pt(13)
                
                # 动作描写（以△开头）
                elif line.startswith('△'):
                    p = self.doc.add_paragraph()
                    p.add_run(line)
                    p.paragraph_format.left_indent = Inches(0.5)
                
                # 角色对话或其他内容
                else:
                    p = self.doc.add_paragraph()
                    p.add_run(line)
            
            # 集尾分隔
            self.doc.add_paragraph('—' * 30)
            self.doc.add_paragraph()
    
    def _is_scene_header(self, line: str) -> bool:
        """
        判断是否是场景头
        支持两种格式：
        - 标准格式：1-1 客厅 日 内
        - Markdown格式：## 1-1 客厅 日 内 或 # 1-1 客厅 日 内
        """
        import re
        # 支持可选的 # 或 ## 前缀
        return bool(re.match(r'^#{0,2}\s*\d+-\d+', line))

