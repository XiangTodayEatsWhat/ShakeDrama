#!/usr/bin/env python3
"""
专业级标准化格式剧本生成的 Web 应用入口。
提供注册、项目管理、配置管理与剧本生成相关 API。
"""
import json
import os
import sys
import threading
import logging
import time
from datetime import datetime
from pathlib import Path
from typing import Optional
from uuid import uuid4

# 项目根目录
ROOT = Path(__file__).resolve().parent
sys.path.insert(0, str(ROOT))

# 加载 .env 文件（如果存在）
DEPLOY_DIR = Path(__file__).resolve().parent
ENV_FILE = DEPLOY_DIR / ".env"
if ENV_FILE.exists():
    with open(ENV_FILE, 'r', encoding='utf-8') as f:
        for line in f:
            line = line.strip()
            # 跳过空行和注释
            if not line or line.startswith('#'):
                continue
            # 解析 KEY=VALUE
            if '=' in line:
                key, value = line.split('=', 1)
                key = key.strip()
                value = value.strip()
                # 只有在环境变量未设置时才设置（允许外部覆盖）
                if key and not os.environ.get(key):
                    os.environ[key] = value
    print(f"已加载环境变量文件: {ENV_FILE}")

PROJECTS_BASE = ROOT / "web_projects"
PROJECTS_BASE.mkdir(exist_ok=True)
OUTPUT_BASE = ROOT / "output"
OUTPUT_BASE.mkdir(exist_ok=True)


def _user_projects_dir(username_key: str) -> Path:
    """返回用户专属的项目目录：web_projects/<username_key>/"""
    d = PROJECTS_BASE / username_key
    d.mkdir(parents=True, exist_ok=True)
    return d


def _user_output_dir(username_key: str, project_id: str) -> Path:
    """返回用户专属的输出目录：output/<username_key>/<project_id>/"""
    d = OUTPUT_BASE / username_key / project_id
    d.mkdir(parents=True, exist_ok=True)
    return d

# 与 web_app 中 stepped_flow 一致
PHASE_IDEA_ENTERED = "idea_entered"

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='[%(asctime)s] [%(levelname)s] %(message)s',
    datefmt='%Y-%m-%d %H:%M:%S'
)

USERS_FILE = ROOT / "users.json"


def _load_json_list(path: Path) -> list:
    if not path.exists():
        return []
    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return data if isinstance(data, list) else []
    except Exception:
        return []


def _save_json_list(path: Path, items: list) -> None:
    with open(path, "w", encoding="utf-8") as f:
        json.dump(items, f, ensure_ascii=False, indent=2)


def load_users() -> list:
    return _load_json_list(USERS_FILE)


def save_users(users: list) -> None:
    _save_json_list(USERS_FILE, users)


def normalize_username(username: str) -> str:
    return "".join((username or "").strip().split()).lower()


def find_user_by_username(username: str) -> Optional[dict]:
    username_key = normalize_username(username)
    if not username_key:
        return None
    return next((u for u in load_users() if u.get("username_key") == username_key), None)


def find_user_by_token(token: str) -> Optional[dict]:
    token = (token or "").strip()
    if not token:
        return None
    return next((u for u in load_users() if u.get("token") == token), None)


def generate_access_token() -> str:
    return f"da_{uuid4().hex}{uuid4().hex}"


def create_user(username: str, password: str = "") -> dict:
    raw_username = (username or "").strip()
    username_key = normalize_username(raw_username)
    if not raw_username:
        raise ValueError("用户名不能为空")
    if len(raw_username) < 2 or len(raw_username) > 32:
        raise ValueError("用户名长度需在 2-32 个字符之间")
    if not password or len(password) < 4:
        raise ValueError("密码不能为空，且至少 4 位")
    if find_user_by_username(raw_username):
        raise ValueError("用户名已注册，请更换后重试")

    import hashlib
    password_hash = hashlib.sha256(password.encode("utf-8")).hexdigest()

    users = load_users()
    user = {
        "username": raw_username,
        "username_key": username_key,
        "password_hash": password_hash,
        "token": generate_access_token(),
        "created_at": datetime.now().isoformat(),
    }
    users.append(user)
    save_users(users)
    return user


def generate_project_id() -> str:
    return str(uuid4())[:8]


def get_project_path(project_id: str, username_key: str = "") -> Path:
    """获取项目 JSON 路径。优先按 username_key 定位；兼容旧数据回退到全局扫描。"""
    if username_key:
        return _user_projects_dir(username_key) / f"{project_id}.json"
    # 兼容：遍历所有用户目录查找
    for user_dir in PROJECTS_BASE.iterdir():
        if user_dir.is_dir():
            p = user_dir / f"{project_id}.json"
            if p.exists():
                return p
    # 兼容旧的根目录存放
    return PROJECTS_BASE / f"{project_id}.json"


def save_project(project: dict) -> None:
    username_key = project.get("owner_username_key", "")
    path = get_project_path(project["id"], username_key)
    path.parent.mkdir(parents=True, exist_ok=True)
    with open(path, "w", encoding="utf-8") as f:
        json.dump(project, f, ensure_ascii=False, indent=2)


def _update_process_storage_final(project: dict, field: str, content: str, version_index: int, user_message: str = ""):
    """更新 process_storage 中某字段的 final_selected"""
    if "process_storage" not in project:
        project["process_storage"] = {}
    if field not in project["process_storage"]:
        project["process_storage"][field] = {"iteration_prompts": [], "final_selected": None}
    project["process_storage"][field]["final_selected"] = {
        "content": content,
        "version_index": version_index,
        "timestamp": datetime.now().isoformat(),
        "user_message": user_message,
    }


def is_invalid_creative_response(text: str) -> bool:
    """检测是否为「非创意内容」的无效回复（如编程助手拒绝、说明文等），避免写入梗概/大纲。"""
    if not (text and text.strip()):
        return True
    t = text.strip()
    invalid_markers = [
        "Claude Code",
        "claude code",
        "不是我的功能范围",
        "不是我的主要功能",
        "命令行工具",
        "专注于软件开发和编程",
        "I'm not designed for creative writing",
        "I'm Claude Code",
        "software development",
        "programming tasks",
        "我需要澄清",
        "由Anthropic开发",
        "核心职能",
        "专注于技术工作",
        "编码助手",
        "编程助手",
        "作为编码助手",
        "作为编程助手",
        "超出了我",
    ]
    return any(m in t for m in invalid_markers)


def create_project(
    user_idea: str,
    total_episodes: int = 80,
    target_audience: str = "通用",
    batch_size: int = 5,
    sample_strategy: str = "auto",
    provider: str = "wlai",
    checkpoint_after_ideation: bool = True,
    username_key: str = "",
) -> dict:
    project_id = generate_project_id()
    
    # 确保输出目录存在（按用户隔离）
    output_dir = _user_output_dir(username_key, project_id) if username_key else _resolve_output_dir(project_id)
    output_dir.mkdir(parents=True, exist_ok=True)
    
    project = {
        "id": project_id,
        "title": "生成中...",
        "user_idea": user_idea.strip(),
        "genre": "",
        "target_audience": target_audience,
        "total_episodes": total_episodes,
        "current_episode": 0,
        "status": "pending",
        "workflow_state": "idle",
        "workflow_phase": PHASE_IDEA_ENTERED,
        "bible_path": str(output_dir / "bible.json"),
        "created_at": datetime.now().isoformat(),
        "updated_at": datetime.now().isoformat(),
        "stepped_flow": True,
        "beat_sheet_chunk_index": 0,
        "script_chunk_end": 0,
        # 新增：分步创意阶段跟踪
        "ideation_stage": "not_started",  # not_started, inspiration, synopsis, characters, overall_outline, multi_outline, beat_sheet, completed
        "inspiration": None,
        # 版本历史管理
        "version_history": {
            "inspiration": [],  # 创作灵感版本历史
            "synopsis": [],     # 故事梗概版本历史
            "overall_outline": [],  # 总体大纲版本历史
            "multi_outline": [],   # 多集大纲版本历史
            "characters": [],   # 人设版本历史
            "beat_sheet": [],   # 分集大纲（多轮对话修改）版本历史
        },
        # 当前选中的版本索引（-1表示最新版本）
        "current_version": {
            "inspiration": -1,
            "synopsis": -1,
            "overall_outline": -1,
            "multi_outline": -1,
            "characters": -1,
            "beat_sheet": -1,
        },
        # 对话历史（按字段分别存储）
        "chat_history": {
            "inspiration": [],
            "synopsis": [],
            "overall_outline": [],
            "multi_outline": [],
            "characters": [],
            "beat_sheet": [],
        },
        # 过程存储：中间用户迭代的 prompt + 最后选定的版本
        "process_storage": {
            "inspiration": {"iteration_prompts": [], "final_selected": None},
            "synopsis": {"iteration_prompts": [], "final_selected": None},
            "overall_outline": {"iteration_prompts": [], "final_selected": None},
            "multi_outline": {"iteration_prompts": [], "final_selected": None},
            "characters": {"iteration_prompts": [], "final_selected": None},
            "beat_sheet": {"iteration_prompts": [], "final_selected": None},
        },
        "config": {
            "batch_size": batch_size,
            "sample_strategy": sample_strategy,
            "provider": provider,
            "checkpoint_after_ideation": checkpoint_after_ideation,
        },
    }
    save_project(project)
    return project


def list_projects(owner_username_key: Optional[str] = None):
    """列出项目；传 owner_username_key 时只返回该用户的项目"""
    projects = []
    if owner_username_key:
        scan_dir = _user_projects_dir(owner_username_key)
    else:
        scan_dir = PROJECTS_BASE
    # 递归扫描 .json（兼容旧的根目录和新的用户子目录）
    for file in scan_dir.rglob("*.json"):
        try:
            with open(file, "r", encoding="utf-8") as f:
                project = json.load(f)
            if owner_username_key and project.get("owner_username_key") != owner_username_key:
                continue
            projects.append(project)
        except Exception:
            continue
    return sorted(projects, key=lambda x: x.get("created_at", ""), reverse=True)


def _get_script_review_range(project: dict) -> tuple[int, int]:
    """获取项目当前剧本审稿范围，优先使用项目已记录的范围。"""
    batch_start = project.get("script_review_batch_start")
    batch_end = project.get("script_review_batch_end")
    if batch_start is not None and batch_end is not None:
        try:
            s = int(batch_start)
            e = int(batch_end)
        except Exception:
            s = int(project.get("current_episode") or 1)
            e = s
        if s > e:
            s, e = e, s
        return s, e

    total = project.get("total_episodes") or project.get("current_episode") or 80
    batch_size = project.get("config", {}).get("batch_size", 5)
    e = int(total)
    s = max(1, e - batch_size + 1)
    return s, e


def _normalize_completed_review_batch(project: dict) -> None:
    """已完成项目：若已有审稿范围则保留，否则回退到最后一批。"""
    if project.get("status") != "completed":
        return
    total = project.get("total_episodes") or project.get("current_episode") or 80
    current = project.get("current_episode", 0)
    if current < total:
        return

    if project.get("script_review_batch_start") is not None and project.get("script_review_batch_end") is not None:
        try:
            s, e = _get_script_review_range(project)
            project["script_review_batch_start"] = s
            project["script_review_batch_end"] = e
            return
        except Exception:
            pass

    s, e = _get_script_review_range({
        "total_episodes": total,
        "current_episode": current,
        "config": project.get("config", {}),
    })
    project["script_review_batch_start"] = s
    project["script_review_batch_end"] = e


def _sync_bible_from_version_history(project: dict, episode_range: Optional[tuple[int, int]] = None) -> int:
    """将 version_history 中已有剧本回填到 bible.json（仅补缺失/空正文）。"""
    version_history = project.get("version_history") or {}
    if not isinstance(version_history, dict) or not version_history:
        return 0

    bible_file = None
    try:
        if project.get("bible_path"):
            bp = Path(project["bible_path"])
            if bp.exists():
                bible_file = bp
    except Exception:
        pass
    if bible_file is None:
        default_bp = _resolve_output_dir(project.get("id", ""), project.get("owner_username_key", "")) / "bible.json"
        if default_bp.exists():
            bible_file = default_bp
    if bible_file is None or not bible_file.exists():
        return 0

    from drama_agent.models.bible import Bible
    from drama_agent.models.episode import Episode

    bible = Bible.load(str(bible_file))
    existing = {ep.number: ep for ep in (bible.episodes or [])}
    current_version = project.get("current_version") or {}

    candidates = []
    for key in version_history.keys():
        if not key.startswith("script_ep_"):
            continue
        try:
            ep_num = int(key.split("script_ep_", 1)[1])
        except Exception:
            continue
        if episode_range is not None:
            start, end = episode_range
            if ep_num < start or ep_num > end:
                continue
        candidates.append(ep_num)

    changed = 0
    for ep_num in sorted(set(candidates)):
        key = f"script_ep_{ep_num}"
        versions = version_history.get(key) or []
        if not versions:
            continue

        chosen = None
        idx = current_version.get(key, -1)
        if isinstance(idx, int) and 0 <= idx < len(versions):
            chosen = versions[idx]
        else:
            chosen = versions[-1]
        content = (chosen.get("content") or "").strip() if isinstance(chosen, dict) else ""
        if not content:
            continue

        if ep_num in existing:
            ep_obj = existing[ep_num]
            if not (ep_obj.full_script or "").strip():
                ep_obj.full_script = content
                changed += 1
            continue

        synopsis = ""
        if bible.beat_sheet and getattr(bible.beat_sheet, "episodes", None):
            beat = next((b for b in bible.beat_sheet.episodes if int(b.get("episode", 0)) == ep_num), None)
            if beat:
                synopsis = (beat.get("synopsis") or "").strip()
        title = (synopsis[:20] if synopsis else "")
        bible.episodes.append(Episode(number=ep_num, title=title, synopsis=synopsis, full_script=content))
        changed += 1

    if changed:
        bible.episodes = sorted(bible.episodes, key=lambda e: e.number)
        if bible.episodes:
            bible.current_episode = max(bible.current_episode or 0, max(ep.number for ep in bible.episodes))
        bible.updated_at = datetime.now().isoformat()
        bible.save(str(bible_file))

    return changed


def _get_memory_baseline_path(project_id: str, start_episode: int, username_key: str = "") -> Path:
    baseline_dir = _resolve_output_dir(project_id, username_key) / "memory_baselines"
    baseline_dir.mkdir(parents=True, exist_ok=True)
    return baseline_dir / f"before_ep_{int(start_episode):04d}.json"


def _save_memory_baseline(project_id: str, start_episode: int, bible) -> tuple[Optional[Path], bool]:
    """保存某批次起始前的 Bible 快照，供人审后回放记忆。返回 (路径, 是否新建)。"""
    if not bible or start_episode is None:
        return None, False

    try:
        start = int(start_episode)
    except Exception:
        return None, False

    if start < 1:
        return None, False

    snapshot_path = _get_memory_baseline_path(project_id, start)
    if snapshot_path.exists():
        return snapshot_path, False

    bible.save(str(snapshot_path))
    return snapshot_path, True


def _load_memory_baseline(project_id: str, start_episode: int):
    """读取某批次起始前的 Bible 快照。"""
    try:
        start = int(start_episode)
    except Exception:
        return None

    snapshot_path = _get_memory_baseline_path(project_id, start)
    if not snapshot_path.exists():
        return None

    from drama_agent.models.bible import Bible

    try:
        return Bible.load(str(snapshot_path))
    except Exception:
        return None


def _replay_review_memory_for_range(project_id: str, workflow, review_start: int, review_end: int) -> dict:
    """按审稿范围重放记忆，返回结构化结果用于日志观测。"""
    result = {
        "replayed": False,
        "used_baseline": False,
        "reviewed_count": 0,
        "start": None,
        "end": None,
        "baseline_path": None,
    }

    if not workflow.context.bible or not workflow.context.bible.episodes:
        return result

    start = int(review_start)
    end = int(review_end)
    if start > end:
        start, end = end, start

    reviewed_episodes = [
        ep for ep in workflow.context.bible.episodes
        if start <= ep.number <= end
    ]
    if not reviewed_episodes:
        result.update({"start": start, "end": end})
        return result

    baseline_path = _get_memory_baseline_path(project_id, start)
    baseline_bible = _load_memory_baseline(project_id, start)
    use_baseline = baseline_bible is not None

    if baseline_bible is not None:
        current_bible = workflow.context.bible
        baseline_bible.episodes = list(current_bible.episodes or [])
        baseline_bible.current_episode = current_bible.current_episode
        baseline_bible.beat_sheet = current_bible.beat_sheet
        replay_base = baseline_bible
    else:
        replay_base = workflow.context.bible

    workflow.context.bible = workflow.memory_manager.run(reviewed_episodes, replay_base)
    result.update({
        "replayed": True,
        "used_baseline": use_baseline,
        "reviewed_count": len(reviewed_episodes),
        "start": start,
        "end": end,
        "baseline_path": str(baseline_path),
    })
    return result


def get_project(project_id: str, username_key: str = ""):
    """获取项目详情"""
    path = get_project_path(project_id, username_key)
    if path.exists():
        with open(path, "r", encoding="utf-8") as f:
            project = json.load(f)
        _normalize_completed_review_batch(project)
        return project
    return None


def _resolve_output_dir(project_id: str, username_key: str = "") -> Path:
    """统一解析项目输出目录。优先按 username_key 定位，否则回退扫描。"""
    if username_key:
        return OUTPUT_BASE / username_key / project_id
    for user_dir in OUTPUT_BASE.iterdir():
        if user_dir.is_dir():
            candidate = user_dir / project_id
            if candidate.is_dir():
                return candidate
    return OUTPUT_BASE / project_id


def get_project_logs(project_id: str, username_key: str = ""):
    """获取项目日志。"""
    import re
    log_file = _resolve_output_dir(project_id, username_key) / "generation.log"
    if not log_file.exists():
        return []
    
    logs = []
    log_re = re.compile(r'^\[([^\]]+)\]\s*\[([^\]]+)\]\s*(.*)$')
    stream_levels = ("STREAM_START", "STREAM_APPEND", "STREAM_APPEND_NL")
    try:
        with open(log_file, "r", encoding="utf-8") as f:
            for line in f:
                line_stripped = line.rstrip("\n\r")
                if not line_stripped:
                    continue
                match = log_re.match(line_stripped)
                if match:
                    ts, level, msg = match.group(1), match.group(2), match.group(3)
                    logs.append({
                        "timestamp": ts,
                        "level": level,
                        "message": msg
                    })
                else:
                    # 不以 [时间] [LEVEL] 开头的行：视为上一行的续行（流式块内换行），合并到上一条
                    if logs and logs[-1]["level"] in stream_levels:
                        logs[-1]["message"] += "\n" + line_stripped
                    else:
                        logs.append({
                            "timestamp": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                            "level": "INFO",
                            "message": line_stripped
                        })
    except Exception as e:
        print(f"读取日志失败: {e}")
    
    return logs


def should_filter_log(message: str) -> bool:
    """判断是否应该过滤此日志（只过滤用户要求的技术细节）"""
    # 精确过滤：只过滤用户明确不想看的内容
    
    # 1. 过滤文件系统路径信息（但保留重要的状态信息）
    if "输出目录:" in message and "/Users/" in message:
        return True
    if "Bible已保存至：" in message and "/Users/" in message:
        return True
    
    # 2. 过滤样本选择的详细信息及趋势缓存提示
    if "已保存趋势数据到缓存" in message:
        return True
    if any(keyword in message for keyword in [
        "选择理由：",   
        "[工作流] 已选择",
        "[工作流] 选择理由",
        "【LLM原始JSON响应】",
        "【LLM原始",
        "```json",
        "```",
        "长度：",
        "字符",
    ]):
        return True
    
    # 3. 过滤HTTP请求日志
    if "INFO:     127.0.0.1" in message or "GET /" in message or "POST /" in message:
        return True
    
    # 4. 过滤空行（但保留有内容的分隔线）
    if message.strip() == "":
        return True
    
    # 其他所有日志都保留（包括Agent消息、进度信息等）
    return False


def write_log(project_id: str, level: str, message: str, newline: bool = True, stream_first: bool = False, username_key: str = ""):
    """写入日志。
    newline=True：正常写一行 [时间] [level] message 并换行。
    newline=False：流式。stream_first=True 时写 [时间] [level] message 不换行；stream_first=False 时只写 message。
    流式结束后调用方要调 end_stream_line(project_id) 补换行。
    若上一笔是流式（未换行），写新一行前会先补换行，保证时间戳单独成行。
    """
    if should_filter_log(message):
        return
    
    log_dir = _resolve_output_dir(project_id, username_key)
    log_dir.mkdir(parents=True, exist_ok=True)
    log_file = log_dir / "generation.log"
    
    need_leading_newline = False
    if newline and log_file.exists() and log_file.stat().st_size > 0:
        try:
            with open(log_file, "rb") as bf:
                bf.seek(-1, 2)
                if bf.read(1) != b"\n":
                    need_leading_newline = True
        except Exception:
            pass
    
    with open(log_file, "a", encoding="utf-8") as f:
        if newline:
            if need_leading_newline:
                f.write("\n")
            timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
            f.write(f"[{timestamp}] [{level}] {message}\n")
        else:
            if stream_first:
                timestamp = datetime.now().strftime('%Y-%m-%d %H:%M:%S')
                f.write(f"[{timestamp}] [{level}] {message}")
            else:
                f.write(message)
        f.flush()


def end_stream_line(project_id: str, username_key: str = ""):
    """流式结束后补一个换行，由调用方在 flush 或切到非流式时调用。"""
    log_file = _resolve_output_dir(project_id, username_key) / "generation.log"
    if not log_file.exists():
        return
    with open(log_file, "a", encoding="utf-8") as f:
        f.write("\n")
        f.flush()


# 线程本地 stdout：多任务并发时，每个线程的 print 只进自己的项目日志
_stdout_wrapper_installed = False
_thread_local_capture = threading.local()


class _ThreadLocalStdoutWrapper:
    """全局唯一 sys.stdout 替换：按线程把 print 转给对应项目的 LogCapture，避免多任务串日志"""
    def __init__(self, real_stdout):
        self._real = real_stdout

    def write(self, message):
        cap = getattr(_thread_local_capture, "current", None)
        if cap is not None:
            cap.write(message)
        else:
            self._real.write(message)

    def flush(self):
        cap = getattr(_thread_local_capture, "current", None)
        if cap is not None:
            cap.flush()
        else:
            self._real.flush()


def _ensure_stdout_wrapper():
    """确保 sys.stdout 已是按线程分发的 wrapper（只安装一次）"""
    global _stdout_wrapper_installed
    if _stdout_wrapper_installed:
        return
    if not isinstance(sys.stdout, _ThreadLocalStdoutWrapper):
        sys.stdout = _ThreadLocalStdoutWrapper(sys.stdout)
        _stdout_wrapper_installed = True


def run_ideation_step(project_id: str, step: str, generate_all: bool = False):
    """执行创意阶段的某一步（分步式）
    
    step可以是：
    - inspiration: 生成创作灵感
    - synopsis: 生成故事梗概
    - characters: 生成人物角色
    - overall_outline: 生成总体大纲
    - multi_outline: 生成多集大纲
    - beat_sheet: 生成分集大纲（generate_all=True 时一次性生成全部，不中途卡点）
    """
    import sys
    
    # 判断是否是「步骤类」日志」（[策划Agent]、[INFO] 等），不是 LLM 正文
    def _is_step_log(line: str) -> bool:
        if not line or len(line) > 500:
            return True
        step_markers = (
            "步骤", "[策划Agent]", "[INFO]", "[SUCCESS]", "[ERROR]", "[WARNING]",
            "爆款", "抓取", "已保存", "趋势", "灵感开始", "灵感生成", "生成故事",
            "生成人物", "生成大纲", "完成", "✅", "❌", "请审核", "点击「继续」",
            "离线已有", "抓取爆款", "第 ", "次抓取", "Bible", "bible",
            # 撰写/审稿阶段：让「第X集：撰写」「正在审核第X集」「审稿完成」等以 INFO 写入并显示
            "[工作流]", "集：撰写", "已入 bible", "一致性检查", "本批第", "审稿", "开始撰写",
            "三角号", "字数与AI", "已通过，进入下一集", "等待人工审稿"
        )
        return any(m in line for m in step_markers)

    class LogCapture:
        """stdout 捕获：流式/非流式、第一块、结束换行 全部由 print 侧通过 stdout_streaming 传参，本类不记状态。"""

        def __init__(self, project_id, echo_stdout):
            self.project_id = project_id
            self.original_stdout = echo_stdout
            self._line_buf = ""  # 非流式时未完成的一行（仅用于按换行切行）

        def write(self, message):
            self.original_stdout.write(message)
            self.original_stdout.flush()

            try:
                from drama_agent.utils.stdout_streaming import (
                    is_stdout_streaming,
                    consume_stdout_stream_first,
                    consume_stdout_stream_end,
                )
                streaming = is_stdout_streaming()
            except Exception:
                streaming = False
                def consume_stdout_stream_first():
                    return False
                def consume_stdout_stream_end():
                    return False
            # 无换行就按流式处理（含长 chunk），避免被当成非流式按行切导致「换行越来越多」
            if not streaming and "\n" not in message:
                streaming = True

            if streaming:
                if not message:
                    return
                # 流式内容里若带换行会变成多行，统一打成空格，避免「后面换行越来越多」
                message = message.replace("\r\n", " ").replace("\n", " ").replace("\r", " ")
                stream_first = consume_stdout_stream_first()
                write_log(
                    self.project_id,
                    "STREAM_START" if stream_first else "STREAM_APPEND",
                    message,
                    newline=False,
                    stream_first=stream_first,
                )
            else:
                if consume_stdout_stream_end():
                    end_stream_line(self.project_id)
                self._line_buf += message
                lines = self._line_buf.split("\n")
                self._line_buf = lines[-1]

                for line in lines[:-1]:
                    if not line.strip():
                        continue
                    if _is_step_log(line):
                        if "ERROR" in line or "错误" in line or "失败" in line:
                            write_log(self.project_id, "ERROR", line)
                        elif "WARNING" in line or "警告" in line:
                            write_log(self.project_id, "WARNING", line)
                        elif "SUCCESS" in line or "成功" in line or "完成" in line or "✅" in line:
                            write_log(self.project_id, "SUCCESS", line)
                        else:
                            write_log(self.project_id, "INFO", line)
                    else:
                        write_log(self.project_id, "STREAM_APPEND_NL", line)

        def flush(self):
            self.original_stdout.flush()
            try:
                from drama_agent.utils.stdout_streaming import consume_stdout_stream_end
                if consume_stdout_stream_end():
                    end_stream_line(self.project_id)
            except Exception:
                pass
            if self._line_buf.strip():
                write_log(self.project_id, "STREAM_APPEND_NL", self._line_buf)
                self._line_buf = ""
    
    log_capture = None
    try:
        project = get_project(project_id)
        if not project:
            write_log(project_id, "ERROR", "项目不存在")
            return
        
        username_key = project.get("owner_username_key", "")
        write_log(project_id, "INFO", f"[步骤] run_ideation_step 开始 step={step}")
        _ensure_stdout_wrapper()
        real_stdout = sys.stdout._real if isinstance(sys.stdout, _ThreadLocalStdoutWrapper) else sys.stdout
        log_capture = LogCapture(project_id, real_stdout)
        _thread_local_capture.current = log_capture
        
        from drama_agent.workflow.drama_workflow import DramaWorkflow
        from drama_agent.sample_library import SelectStrategy
        from drama_agent.config import Config
        from drama_agent.agents.showrunner import ShowrunnerAgent
        
        output_dir = _resolve_output_dir(project_id, username_key)
        bible_path = output_dir / "bible.json"
        
        config = Config()
        config.output_dir = str(output_dir)
        config.bible_path = str(bible_path)
        
        from drama_agent.config import resolve_provider_to_config, set_config
        from drama_agent.utils.llm_client import reset_llm_client
        provider = project.get("config", {}).get("provider", "wlai")
        resolve_provider_to_config(config, provider)
        set_config(config)
        reset_llm_client()

        showrunner = ShowrunnerAgent()
        showrunner.config = config
        
        # 加载或创建Bible
        from drama_agent.models import Bible
        if bible_path.exists():
            with open(bible_path, 'r', encoding='utf-8') as f:
                import json
                bible_data = json.load(f)
                bible = Bible.from_dict(bible_data)
        else:
            bible = Bible(
                title="生成中...",
                genre=[],
                target_audience=project.get("target_audience", "通用"),
                synopsis="",
                total_episodes=project.get("total_episodes", 80)
            )
        
        # 根据不同的step执行不同的操作
        if step == "inspiration":
            write_log(project_id, "INFO", "=" * 60)
            write_log(project_id, "INFO", "[步骤 1/6] 生成创作灵感")
            write_log(project_id, "INFO", "=" * 60)
            
            # 获取趋势
            try:
                write_log(project_id, "INFO", "🔥 爆款抓取开始...")
                from drama_agent.utils.trend_search import search_short_drama_trends
                trend_hint = search_short_drama_trends(debug=False)
                write_log(project_id, "SUCCESS", "✅ 爆款抓取结束")
            except Exception as e:
                write_log(project_id, "WARNING", f"⚠️ 爆款抓取失败: {e}")
                trend_hint = ""
            
            write_log(project_id, "INFO", "[策划Agent] 💡 灵感开始生成...")
            write_log(project_id, "INFO", "=" * 60)
            
            # 生成灵感（会实时流式输出到stdout）
            inspiration = showrunner.generate_inspiration(
                project.get("user_idea", ""),
                reference_style=None,
                trend_hint=trend_hint
            )
            
            write_log(project_id, "INFO", "=" * 60)
            write_log(project_id, "SUCCESS", "[策划Agent] ✅ 灵感生成结束")
            
            # 保存灵感到项目
            project["inspiration"] = inspiration
            project["ideation_stage"] = "inspiration"
            project["status"] = "checkpoint"
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
            
            write_log(project_id, "SUCCESS", "=" * 60)
            write_log(project_id, "SUCCESS", "✅ 创作灵感已生成！")
            write_log(project_id, "INFO", "📋 请审核创作灵感")
            write_log(project_id, "INFO", "👉 审核通过后点击「继续」生成故事梗概")
            write_log(project_id, "SUCCESS", "=" * 60)
            
        elif step == "synopsis":
            write_log(project_id, "INFO", "=" * 60)
            write_log(project_id, "INFO", "[步骤 2/6] 生成故事梗概")
            write_log(project_id, "INFO", "=" * 60)
            
            inspiration = project.get("inspiration", "")
            
            write_log(project_id, "INFO", "[策划Agent] 正在生成故事梗概...")
            synopsis_data = showrunner.generate_synopsis(
                project.get("user_idea", ""),
                reference_style=None,
                inspiration=inspiration
            )
            
            total_episodes = project.get("total_episodes", 80)
            if total_episodes:
                synopsis_data["total_episodes"] = total_episodes
            
            write_log(project_id, "SUCCESS", "[策划Agent] ✅ 故事梗概生成完成")
            write_log(project_id, "INFO", f"[策划Agent] 📖 剧本标题：{synopsis_data['title']}")
            write_log(project_id, "INFO", f"[策划Agent] 🎭 故事类型：{synopsis_data['genre']}")
            write_log(project_id, "INFO", f"[策划Agent] 📝 梗概：\n{synopsis_data['synopsis']}")
            
            # 更新Bible
            bible.title = synopsis_data["title"]
            bible.genre = synopsis_data["genre"]
            bible.synopsis = synopsis_data["synopsis"]
            bible.theme = synopsis_data.get("theme", "")
            bible.total_episodes = synopsis_data.get("total_episodes", 80)
            bible.save(str(bible_path))
            
            # 更新项目
            project["title"] = synopsis_data["title"]
            project["genre"] = synopsis_data["genre"]
            project["ideation_stage"] = "synopsis"
            project["status"] = "checkpoint"
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
            
            write_log(project_id, "SUCCESS", "=" * 60)
            write_log(project_id, "SUCCESS", "✅ 故事梗概已生成！")
            write_log(project_id, "INFO", "📋 请审核故事梗概")
            write_log(project_id, "INFO", "👉 审核通过后点击「继续」生成人物角色")
            write_log(project_id, "SUCCESS", "=" * 60)
            
        elif step == "characters":
            write_log(project_id, "INFO", "=" * 60)
            write_log(project_id, "INFO", "[步骤 3/6] 生成人物角色")
            write_log(project_id, "INFO", "=" * 60)
            
            inspiration = project.get("inspiration", "")
            
            write_log(project_id, "INFO", "[策划Agent] 正在创建人物角色...")
            characters_data = showrunner.create_characters(
                bible.synopsis,
                bible.genre,
                bible.target_audience,
                overall_outline="",
                inspiration=inspiration
            )
            write_log(project_id, "SUCCESS", f"[策划Agent] ✅ 人物角色创建完成（共{len(characters_data['characters'])}位）")
            
            # 添加角色到Bible（characters 为 Dict[str, Character]，不能用 list）
            from drama_agent.models.character import Character, CharacterArchetype, CharacterRelationship
            bible.characters = {}
            for char_data in characters_data["characters"]:
                age_val = char_data.get("age")
                if isinstance(age_val, str):
                    import re
                    m = re.search(r"\d+", age_val)
                    age_val = int(m.group(0)) if m else None
                
                bg = char_data.get("background", "")
                if char_data.get("core_goal"):
                    bg += " 核心目标：" + str(char_data["core_goal"])
                if char_data.get("memory_point"):
                    bg += " 记忆点：" + str(char_data["memory_point"])
                
                character = Character(
                    name=char_data["name"],
                    identity=char_data["identity"],
                    archetype=CharacterArchetype(char_data["archetype"]),
                    age=age_val,
                    personality=char_data.get("personality", ""),
                    background=bg,
                    skills=char_data.get("skills", []),
                    secrets=char_data.get("secrets", []),
                    arc=char_data.get("arc", "")
                )
                bible.add_character(character)
                
                if char_data["archetype"] == "protagonist":
                    bible.protagonist_name = char_data["name"]
                
                write_log(project_id, "INFO", f"[策划Agent] 👤 {char_data['name']} - {char_data['identity']}")
            
            # 添加角色关系
            for rel in characters_data.get("relationships", []):
                char1 = bible.get_character(rel["character1"])
                if char1:
                    char1.add_relationship(CharacterRelationship(
                        target=rel["character2"],
                        relation_type=rel["relation_type"],
                        sentiment="neutral",
                        notes=rel.get("dynamic", "")
                    ))
            
            bible.save(str(bible_path))
            
            project["ideation_stage"] = "characters"
            project["status"] = "checkpoint"
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
            
            write_log(project_id, "SUCCESS", "=" * 60)
            write_log(project_id, "SUCCESS", "✅ 人物角色已创建！")
            write_log(project_id, "INFO", "📋 请审核人物角色设定")
            write_log(project_id, "INFO", "👉 审核通过后点击「继续」生成总体大纲")
            write_log(project_id, "SUCCESS", "=" * 60)
            
        elif step == "overall_outline":
            write_log(project_id, "INFO", "=" * 60)
            write_log(project_id, "INFO", "[步骤 4/6] 生成总体大纲")
            write_log(project_id, "INFO", "=" * 60)
            
            inspiration = project.get("inspiration", "")
            
            write_log(project_id, "INFO", "[策划Agent] 正在生成总体大纲...")
            overall_outline = showrunner.generate_overall_outline(
                bible.synopsis,
                bible.title,
                bible.genre,
                total_episodes=bible.total_episodes,
                inspiration=inspiration,
                characters=[{
                    "name": c.name,
                    "identity": c.identity,
                    "archetype": c.archetype.value
                } for c in bible.characters.values()]
            )
            write_log(project_id, "SUCCESS", "[策划Agent] ✅ 总体大纲生成完成")
            write_log(project_id, "INFO", f"[策划Agent] 📋 总体大纲：\n{overall_outline}")
            
            bible.overall_outline = overall_outline
            bible.save(str(bible_path))
            
            project["ideation_stage"] = "overall_outline"
            project["status"] = "checkpoint"
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
            
            write_log(project_id, "SUCCESS", "=" * 60)
            write_log(project_id, "SUCCESS", "✅ 总体大纲已生成！")
            write_log(project_id, "INFO", "📋 请审核总体大纲")
            write_log(project_id, "INFO", "👉 审核通过后点击「继续」生成多集大纲")
            write_log(project_id, "SUCCESS", "=" * 60)
            
        elif step == "multi_outline":
            write_log(project_id, "INFO", "=" * 60)
            write_log(project_id, "INFO", "[步骤 5/6] 生成多集大纲")
            write_log(project_id, "INFO", "=" * 60)
            
            inspiration = project.get("inspiration", "")
            write_log(project_id, "INFO", f"[策划Agent] 正在生成{bible.total_episodes}集的多集大纲...")
            multi_episode_outline = showrunner.generate_multi_episode_outline(
                bible.synopsis,
                bible.overall_outline,
                [{
                    "name": c.name,
                    "identity": c.identity,
                    "archetype": c.archetype.value
                } for c in bible.characters.values()],
                total_episodes=bible.total_episodes,
                inspiration=inspiration
            )
            write_log(project_id, "SUCCESS", "[策划Agent] ✅ 多集大纲生成完成")
            write_log(project_id, "INFO", f"[策划Agent] 📋 多集大纲：\n{multi_episode_outline}")
            
            bible.multi_episode_outline = multi_episode_outline
            bible.save(str(bible_path))
            
            project["ideation_stage"] = "multi_outline"
            project["status"] = "checkpoint"
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
            
            write_log(project_id, "SUCCESS", "=" * 60)
            write_log(project_id, "SUCCESS", "✅ 多集大纲已生成！")
            write_log(project_id, "INFO", "📋 请审核多集大纲")
            write_log(project_id, "INFO", "👉 审核通过后点击「继续」生成分集大纲")
            write_log(project_id, "SUCCESS", "=" * 60)
            
        elif step == "beat_sheet":
            total_episodes = bible.total_episodes
            from drama_agent.models.episode import BeatSheet
            inspiration = project.get("inspiration", "")
            characters_list = [{"name": c.name, "identity": c.identity, "archetype": c.archetype.value} for c in bible.characters.values()]

            while True:
                next_start = project.get("beat_sheet_next_start", 1)
                batch_end = min(next_start + 9, total_episodes)

                write_log(project_id, "INFO", "=" * 60)
                if generate_all:
                    write_log(project_id, "INFO", f"[步骤 6/6] 一次性生成分集大纲（第{next_start}-{batch_end}集，目标全部{total_episodes}集）")
                else:
                    write_log(project_id, "INFO", f"[步骤 6/6] 生成分集大纲（第{next_start}-{batch_end}集）")
                write_log(project_id, "INFO", "=" * 60)

                existing_beats = []
                if bible.beat_sheet and bible.beat_sheet.episodes:
                    existing_beats = [
                        {"episode": b["episode"], "synopsis": b["synopsis"], "ending_hook": b["ending_hook"], "hook_type": b.get("hook_type", "cliffhanger")}
                        for b in bible.beat_sheet.episodes
                    ]

                write_log(project_id, "INFO", f"[策划Agent] 正在生成第{next_start}-{batch_end}集...")
                beat_sheet_data = showrunner.generate_beat_sheet_batch(
                    next_start,
                    batch_end,
                    total_episodes,
                    bible.synopsis,
                    characters_list,
                    existing_beats,
                    multi_episode_outline=bible.multi_episode_outline,
                    inspiration=inspiration
                )
                write_log(project_id, "SUCCESS", f"[策划Agent] ✅ 第{next_start}-{batch_end}集分集大纲生成完成（共{len(beat_sheet_data['beats'])}集）")

                if not bible.beat_sheet:
                    bible.beat_sheet = BeatSheet()
                for beat in beat_sheet_data["beats"]:
                    bible.beat_sheet.add_beat(
                        beat["episode"],
                        beat["synopsis"],
                        beat["ending_hook"],
                        beat.get("hook_type", "cliffhanger")
                    )
                if next_start == 1 and beat_sheet_data.get("beats"):
                    for beat in beat_sheet_data["beats"][:5]:
                        key_conflict = beat.get("key_conflict")
                        if key_conflict:
                            bible.active_conflicts.append(key_conflict)

                bible.save(str(bible_path))

                if batch_end >= total_episodes:
                    project["ideation_stage"] = "completed"
                    project["status"] = "checkpoint"
                    project.pop("beat_sheet_next_start", None)
                    project["updated_at"] = datetime.now().isoformat()
                    write_log(project_id, "SUCCESS", "=" * 60)
                    write_log(project_id, "SUCCESS", "✅ 创意策划全部完成！")
                    write_log(project_id, "INFO", f"[策划Agent] 🎉 《{bible.title}》创意策划完成")
                    write_log(project_id, "INFO", f"[策划Agent] 📊 已规划{total_episodes}集，创建{len(bible.characters)}位角色")
                    write_log(project_id, "INFO", "📋 请审核完整的创意策划内容")
                    write_log(project_id, "INFO", "👉 审核通过后点击「开始撰写剧本」")
                    write_log(project_id, "SUCCESS", "=" * 60)
                    save_project(project)
                    break
                else:
                    next_batch_start = batch_end + 1
                    next_batch_end = min(next_batch_start + 9, total_episodes)
                    project["beat_sheet_next_start"] = next_batch_start
                    project["updated_at"] = datetime.now().isoformat()

                    if generate_all:
                        write_log(project_id, "SUCCESS", f"✅ 第{next_start}-{batch_end}集完成，继续生成第{next_batch_start}-{next_batch_end}集...")
                        save_project(project)
                        continue
                    else:
                        project["ideation_stage"] = "beat_sheet"
                        project["status"] = "checkpoint"
                        write_log(project_id, "SUCCESS", "=" * 60)
                        write_log(project_id, "SUCCESS", f"✅ 第{next_start}-{batch_end}集分集大纲已生成！")
                        write_log(project_id, "INFO", "📋 请审核本批分集大纲")
                        write_log(project_id, "INFO", f"👉 审核通过后点击「继续」生成第{next_batch_start}-{next_batch_end}集")
                        write_log(project_id, "SUCCESS", "=" * 60)
                        save_project(project)
                        break
        
    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        write_log(project_id, "ERROR", f"执行失败: {str(e)}")
        write_log(project_id, "ERROR", error_detail)
        
        project = get_project(project_id)
        if project:
            project["status"] = "failed"
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
    finally:
        # 写出 line_buffer 中可能存在的未完成行（实践中多为 ''）；按线程清除
        if log_capture is not None:
            try:
                log_capture.flush()
            except Exception:
                pass
        _thread_local_capture.current = None


def run_generation_task(project_id: str):
    """后台运行生成任务"""
    import sys
    from io import StringIO
    
    # 重定向标准输出到日志（线程本地，与 run_ideation_step 一致）
    class LogCapture:
        def __init__(self, project_id, echo_stdout):
            self.project_id = project_id
            self.original_stdout = echo_stdout

        def write(self, message):
            self.original_stdout.write(message)
            self.original_stdout.flush()

            if message.strip():
                if "ERROR" in message or "错误" in message or "失败" in message:
                    level = "ERROR"
                elif "WARNING" in message or "警告" in message:
                    level = "WARNING"
                elif "SUCCESS" in message or "成功" in message or "完成" in message or "✅" in message:
                    level = "SUCCESS"
                else:
                    level = "INFO"
                write_log(self.project_id, level, message.strip())

        def flush(self):
            self.original_stdout.flush()

    try:
        write_log(project_id, "INFO", "=" * 60)
        write_log(project_id, "INFO", "🎬 开始生成任务")
        write_log(project_id, "INFO", "=" * 60)

        project = get_project(project_id)
        if not project:
            write_log(project_id, "ERROR", "项目不存在")
            return

        username_key = project.get("owner_username_key", "")
        _ensure_stdout_wrapper()
        real_stdout = sys.stdout._real if isinstance(sys.stdout, _ThreadLocalStdoutWrapper) else sys.stdout
        log_capture = LogCapture(project_id, real_stdout)
        _thread_local_capture.current = log_capture

        # 导入workflow
        from drama_agent.workflow.drama_workflow import DramaWorkflow
        from drama_agent.sample_library import SelectStrategy
        from drama_agent.config import Config
        
        # 配置输出目录
        output_dir = _resolve_output_dir(project_id, username_key)
        output_dir.mkdir(parents=True, exist_ok=True)
        
        # 创建配置
        config = Config()
        config.output_dir = str(output_dir)
        config.bible_path = str(output_dir / "bible.json")
        config.drama.episodes_per_batch = project.get("config", {}).get("batch_size", 5)
        
        # 设置provider（支持自定义模型）
        from drama_agent.config import resolve_provider_to_config, set_config
        from drama_agent.utils.llm_client import reset_llm_client
        provider = project.get("config", {}).get("provider", "wlai")
        resolve_provider_to_config(config, provider)
        set_config(config)
        reset_llm_client()

        write_log(project_id, "INFO", f"使用AI模型: {config.llm.provider}")
        write_log(project_id, "INFO", f"输出目录: {output_dir}")
        write_log(project_id, "INFO", f"目标生成: {project.get('total_episodes', 80)} 集")
        write_log(project_id, "INFO", f"批次大小: {config.drama.episodes_per_batch} 集")
        
        # 初始化workflow
        write_log(project_id, "INFO", "正在初始化生成引擎...")
        workflow = DramaWorkflow()
        workflow.config = config
        
        # 设置日志回调
        original_log = workflow._log
        def custom_log(message: str):
            original_log(message)
            # 解析日志级别
            if "错误" in message or "失败" in message:
                level = "ERROR"
            elif "警告" in message:
                level = "WARNING"
            elif "完成" in message or "成功" in message:
                level = "SUCCESS"
            else:
                level = "INFO"
            write_log(project_id, level, message)
            
            # 实时更新进度和标题
            try:
                current_project = get_project(project_id)
                if not current_project:
                    return
                
                # 更新标题（从bible获取）
                if workflow.context.bible and workflow.context.bible.title:
                    if current_project.get("title") != workflow.context.bible.title:
                        current_project["title"] = workflow.context.bible.title
                        write_log(project_id, "INFO", f"标题更新为: {workflow.context.bible.title}")
                
                # 更新进度
                if workflow.context.bible and workflow.context.bible.episodes:
                    current_project["current_episode"] = len(workflow.context.bible.episodes)
                
                current_project["updated_at"] = datetime.now().isoformat()
                save_project(current_project)
            except Exception:
                pass  # 忽略进度更新错误
        
        workflow._log = custom_log
        
        # 获取样本策略
        sample_strategy_str = project.get("config", {}).get("sample_strategy", "auto")
        sample_strategy = SelectStrategy.AUTO if sample_strategy_str == "auto" else SelectStrategy.NONE
        
        write_log(project_id, "INFO", "=" * 60)
        write_log(project_id, "INFO", "开始创意策划阶段")
        write_log(project_id, "INFO", "=" * 60)
        
        # 检查是否需要在创意阶段后卡点
        checkpoint_enabled = project.get("config", {}).get("checkpoint_after_ideation", True)
        
        if checkpoint_enabled:
            # 只执行创意阶段 - 分步执行以实现在灵感生成后停止
            write_log(project_id, "INFO", "🎯 启用卡点模式：将在创意策划完成后暂停")
            
            # 执行创意阶段（run_ideation内部会完成所有创意工作）
            bible = workflow.run_ideation(
                user_idea=project.get("user_idea", ""),
                sample_strategy=sample_strategy,
                manual_sample_ids=None,
                total_episodes=project.get("total_episodes", 80)
            )
            
            # 更新项目信息
            project["title"] = bible.title if bible else "未命名项目"
            project["status"] = "checkpoint"  # 卡点状态
            project["current_episode"] = 0
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
            
            write_log(project_id, "SUCCESS", "=" * 60)
            write_log(project_id, "SUCCESS", "✅ 创意策划完成！")
            write_log(project_id, "INFO", "📋 请在详情页审核创意内容")
            write_log(project_id, "INFO", f"📝 剧本标题：{bible.title}")
            write_log(project_id, "INFO", f"📚 故事类型：{bible.genre}")
            write_log(project_id, "INFO", f"👥 主要角色：{len(bible.characters)}位")
            write_log(project_id, "INFO", "👉 审核通过后点击「继续生成」按钮开始撰写剧本")
            write_log(project_id, "SUCCESS", "=" * 60)
            
        else:
            # 运行完整流程
            write_log(project_id, "INFO", "⚡ 快速模式：直接生成完整剧本")
            bible = workflow.run_full(
                user_idea=project.get("user_idea", ""),
                sample_strategy=sample_strategy,
                episodes_per_batch=config.drama.episodes_per_batch,
                total_episodes=project.get("total_episodes", 80)
            )
            
            # 更新项目信息
            project["title"] = bible.title if bible else "未命名项目"
            project["status"] = "completed"
            project["current_episode"] = len(bible.episodes) if bible and bible.episodes else 0
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
        
        write_log(project_id, "SUCCESS", "=" * 60)
        write_log(project_id, "SUCCESS", f"🎉 生成任务完成！共生成 {project['current_episode']} 集")
        write_log(project_id, "SUCCESS", "=" * 60)

    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        write_log(project_id, "ERROR", f"生成失败: {str(e)}")
        write_log(project_id, "ERROR", f"详细错误信息已记录")

        project = get_project(project_id)
        if project:
            project["status"] = "failed"
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
    finally:
        _thread_local_capture.current = None


def run_scripting_task(project_id: str, script_review_batch_start=None, script_review_batch_end=None, auto_continue_to: int = 0, pending_review_start=None):
    """继续生成剧本（从卡点恢复）。传入上批审稿范围后先做记忆回放，再撰写下一批；auto_continue_to>0 时可连续写到目标集。"""
    import sys

    class LogCapture:
        def __init__(self, project_id, echo_stdout):
            self.project_id = project_id
            self.original_stdout = echo_stdout

        def write(self, message):
            self.original_stdout.write(message)
            self.original_stdout.flush()
            if message.strip():
                if "ERROR" in message or "错误" in message or "失败" in message:
                    level = "ERROR"
                elif "WARNING" in message or "警告" in message:
                    level = "WARNING"
                elif "SUCCESS" in message or "成功" in message or "完成" in message or "✅" in message:
                    level = "SUCCESS"
                else:
                    level = "INFO"
                write_log(self.project_id, level, message.strip())

        def flush(self):
            self.original_stdout.flush()

    log_capture = None
    try:
        write_log(project_id, "INFO", "=" * 60)
        write_log(project_id, "INFO", "📝 继续生成剧本")
        write_log(project_id, "INFO", "=" * 60)

        project = get_project(project_id)
        if not project:
            write_log(project_id, "ERROR", "项目不存在")
            return

        username_key = project.get("owner_username_key", "")
        _ensure_stdout_wrapper()
        real_stdout = sys.stdout._real if isinstance(sys.stdout, _ThreadLocalStdoutWrapper) else sys.stdout
        log_capture = LogCapture(project_id, real_stdout)
        _thread_local_capture.current = log_capture

        # 导入workflow
        from drama_agent.workflow.drama_workflow import DramaWorkflow
        from drama_agent.config import Config

        output_dir = _resolve_output_dir(project_id, username_key)
        bible_path = output_dir / "bible.json"

        if not bible_path.exists():
            write_log(project_id, "ERROR", "未找到创意策划结果，请重新生成")
            project["status"] = "failed"
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
            return
        
        # 创建配置
        config = Config()
        config.output_dir = str(output_dir)
        config.bible_path = str(bible_path)
        config.drama.episodes_per_batch = project.get("config", {}).get("batch_size", 5)
        
        # 设置provider（支持自定义模型）
        from drama_agent.config import resolve_provider_to_config, set_config
        from drama_agent.utils.llm_client import reset_llm_client
        provider = project.get("config", {}).get("provider", "wlai")
        resolve_provider_to_config(config, provider)
        set_config(config)
        reset_llm_client()

        write_log(project_id, "INFO", f"使用AI模型: {config.llm.provider}")
        write_log(project_id, "INFO", f"批次大小: {config.drama.episodes_per_batch} 集")
        
        # 初始化 workflow 并从 bible 恢复进度（支持失败后从未完成的集数继续）
        write_log(project_id, "INFO", "正在加载创意策划结果...")
        workflow = DramaWorkflow()
        workflow.config = config
        workflow.resume(str(bible_path))
        
        # 设置日志回调
        original_log = workflow._log
        def custom_log(message: str):
            original_log(message)
            # 解析日志级别
            if "错误" in message or "失败" in message:
                level = "ERROR"
            elif "警告" in message:
                level = "WARNING"
            elif "完成" in message or "成功" in message:
                level = "SUCCESS"
            else:
                level = "INFO"
            write_log(project_id, level, message)
            
            # 实时更新进度
            try:
                current_project = get_project(project_id)
                if not current_project:
                    return
                
                # 更新进度
                if workflow.context.bible and workflow.context.bible.episodes:
                    current_project["current_episode"] = len(workflow.context.bible.episodes)
                
                current_project["updated_at"] = datetime.now().isoformat()
                save_project(current_project)
            except Exception:
                pass  # 忽略进度更新错误
        
        workflow._log = custom_log

        # 人审范围（由 continue / continue_to 传入；若缺失则回退到项目当前记录）
        r_start = script_review_batch_start
        r_end = script_review_batch_end
        if r_start is None or r_end is None:
            project_for_memory = get_project(project_id)
            if project_for_memory:
                r_start = r_start if r_start is not None else project_for_memory.get("script_review_batch_start")
                r_end = r_end if r_end is not None else project_for_memory.get("script_review_batch_end")

        review_anchor_start = pending_review_start
        if auto_continue_to > 0 and review_anchor_start is None:
            try:
                if r_end is not None:
                    review_anchor_start = int(r_end) + 1
                elif r_start is not None:
                    review_anchor_start = int(r_start)
            except Exception:
                review_anchor_start = None

        # 人审通过后：按审稿范围回放记忆。若命中基线快照则从“审稿起点前”重算该段记忆。
        if r_start is not None and r_end is not None:
            try:
                replay_start = int(r_start)
                replay_end = int(r_end)
                if replay_start > replay_end:
                    replay_start, replay_end = replay_end, replay_start

                write_log(project_id, "INFO", f"人审已通过，开始回放记忆：第{replay_start}-{replay_end}集")
                replay_t0 = time.perf_counter()
                replay_result = _replay_review_memory_for_range(
                    project_id,
                    workflow,
                    replay_start,
                    replay_end,
                )
                replay_elapsed_ms = int((time.perf_counter() - replay_t0) * 1000)

                if replay_result["replayed"]:
                    if workflow.context.bible:
                        workflow.context.bible.save(str(bible_path))
                    mode_text = "快照重算" if replay_result["used_baseline"] else "增量回放"
                    write_log(
                        project_id,
                        "INFO",
                        f"记忆回放完成：模式={mode_text}，范围第{replay_result['start']}-{replay_result['end']}集，"
                        f"处理{replay_result['reviewed_count']}集，耗时{replay_elapsed_ms}ms，"
                        f"快照={replay_result['baseline_path'] or '无'}"
                    )
                else:
                    write_log(
                        project_id,
                        "INFO",
                        f"记忆回放跳过：范围第{replay_start}-{replay_end}集无可处理剧本，耗时{replay_elapsed_ms}ms"
                    )
            except Exception as mem_err:
                write_log(project_id, "WARNING", f"记忆官回放失败（继续撰写）: {mem_err}")

        if auto_continue_to > 0:
            target_info = f"，目标第{auto_continue_to}集" if auto_continue_to else ""
            write_log(project_id, "INFO", f"连续生成模式已启用{target_info}")

        # 为下一批起始集保存基线快照，供后续人审通过后做“起点前重算”
        try:
            next_batch_start = getattr(workflow.context, "current_batch_start", 0)
            if workflow.context.bible and next_batch_start:
                baseline_path, created = _save_memory_baseline(project_id, int(next_batch_start), workflow.context.bible)
                if baseline_path:
                    baseline_status = "新建" if created else "复用"
                    write_log(project_id, "INFO", f"记忆基线快照已{baseline_status}：第{int(next_batch_start)}集起点 -> {baseline_path}")
        except Exception as baseline_err:
            write_log(project_id, "WARNING", f"记忆基线快照保存失败（不影响继续）: {baseline_err}")
        
        write_log(project_id, "INFO", "=" * 60)
        write_log(project_id, "INFO", "开始撰写剧本")
        write_log(project_id, "INFO", "=" * 60)
        
        def _on_episode_done(ep_num: int, ep):
            """每集完成后追加一条剧本版本（可选用中间版本、保存并继续）"""
            p = get_project(project_id)
            if not p:
                return
            if "version_history" not in p:
                p["version_history"] = {}
            key = f"script_ep_{ep_num}"
            if key not in p["version_history"]:
                p["version_history"][key] = []
            vers = p["version_history"][key]
            parent = len(vers) - 1 if vers else -1
            vers.append({
                "content": ep.full_script or "",
                "timestamp": datetime.now().isoformat(),
                "user_message": "",
                "ai_response": "",
                "version_number": len(vers) + 1,
                "parent_version_index": parent,
            })
            p["updated_at"] = datetime.now().isoformat()
            save_project(p)
        
        # 一集一集：写→审→不通过则多轮重写该集→通过后入 bible；每集落版入 version_history
        from drama_agent.workflow.drama_workflow import WorkflowState
        workflow.run_scripting_batch(
            format_only_review=True,
            stop_after_batch_for_human=True,
            on_episode_done=_on_episode_done,
        )
        
        if workflow.context.state == WorkflowState.ERROR:
            write_log(project_id, "ERROR", f"错误：{workflow.context.error_message}")
            if workflow.context.bible:
                try:
                    workflow.context.bible.save(str(bible_path))
                except Exception:
                    pass
            project = get_project(project_id)
            if project:
                project["status"] = "failed"
                project["updated_at"] = datetime.now().isoformat()
                save_project(project)
            return
        
        # 本批成功：落盘保存
        if workflow.context.bible:
            try:
                workflow.context.bible.save(str(bible_path))
            except Exception as save_err:
                write_log(project_id, "WARNING", f"保存进度失败（不影响继续）: {save_err}")
        
        project = get_project(project_id)
        if not project:
            return
        project["current_episode"] = len(workflow.context.bible.episodes) if workflow.context.bible and workflow.context.bible.episodes else 0
        project["updated_at"] = datetime.now().isoformat()
        
        if workflow.context.state == WorkflowState.SCRIPT_REVIEW:
            r_start = getattr(workflow.context, "review_batch_start", 0) or workflow.context.current_batch_start
            r_end = getattr(workflow.context, "review_batch_end", 0) or workflow.context.current_batch_end
            batch_size = project.get("config", {}).get("batch_size", 5)
            if r_start > r_end:
                r_start = max(1, r_end - batch_size + 1)
            try:
                _sync_bible_from_version_history(project, (int(r_start), int(r_end)))
            except Exception:
                pass

            # auto_continue_to: 跳过审稿，连续生成到目标集数
            if auto_continue_to > 0 and r_end < auto_continue_to:
                write_log(project_id, "INFO", f"连续生成模式：第{r_start}-{r_end}集完成，继续生成...")
                if "version_history" not in project:
                    project["version_history"] = {}
                if "current_version" not in project:
                    project["current_version"] = {}
                bible = workflow.context.bible
                if bible and bible.episodes:
                    for ep_num in range(r_start, r_end + 1):
                        key = f"script_ep_{ep_num}"
                        if key not in project["version_history"]:
                            project["version_history"][key] = []
                        vers = project["version_history"][key]
                        if len(vers) > 1:
                            latest = dict(vers[-1])
                            latest["version_number"] = 1
                            latest["parent_version_index"] = -1
                            project["version_history"][key] = [latest]
                            project["current_version"][key] = 0
                            _update_process_storage_final(project, key, latest.get("content", ""), 0, latest.get("user_message", ""))
                        elif len(vers) == 1:
                            vers[0]["version_number"] = 1
                            vers[0]["parent_version_index"] = -1
                            project["current_version"][key] = 0
                            _update_process_storage_final(project, key, vers[0].get("content", ""), 0, vers[0].get("user_message", ""))
                        if len(project["version_history"][key]) == 0:
                            ep = next((e for e in bible.episodes if e.number == ep_num), None)
                            if ep and (ep.full_script or "").strip():
                                c = (ep.full_script or "").strip()
                                project["version_history"][key].append({
                                    "content": c,
                                    "timestamp": datetime.now().isoformat(),
                                    "user_message": "",
                                    "ai_response": "",
                                    "version_number": 1,
                                    "parent_version_index": -1,
                                })
                                project["current_version"][key] = 0
                                _update_process_storage_final(project, key, c, 0, "")
                save_project(project)
                workflow.run_scripting_batch(
                    format_only_review=True,
                    stop_after_batch_for_human=True,
                    on_episode_done=_on_episode_done,
                )
                if workflow.context.bible:
                    try:
                        workflow.context.bible.save(str(bible_path))
                    except Exception as save_err:
                        write_log(project_id, "WARNING", "保存进度失败（不影响继续）: " + str(save_err))
                project = get_project(project_id)
                if not project:
                    return
                project["current_episode"] = len(workflow.context.bible.episodes) if workflow.context.bible and workflow.context.bible.episodes else 0
                project["updated_at"] = datetime.now().isoformat()
                if workflow.context.state == WorkflowState.SCRIPT_REVIEW:
                    r_start2 = getattr(workflow.context, "review_batch_start", 0) or workflow.context.current_batch_start
                    r_end2 = getattr(workflow.context, "review_batch_end", 0) or workflow.context.current_batch_end
                    if r_start2 > r_end2:
                        r_start2 = max(1, r_end2 - batch_size + 1)
                    if r_end2 < auto_continue_to:
                        save_project(project)
                        return run_scripting_task(
                            project_id,
                            script_review_batch_start=r_start2,
                            script_review_batch_end=r_end2,
                            auto_continue_to=auto_continue_to,
                            pending_review_start=review_anchor_start,
                        )
                    else:
                        try:
                            _sync_bible_from_version_history(project, (int(r_start2), int(r_end2)))
                        except Exception:
                            pass
                        project["status"] = "checkpoint"
                        review_start = review_anchor_start if review_anchor_start is not None else r_start2
                        project["script_review_batch_start"] = review_start
                        project["script_review_batch_end"] = r_end2
                        save_project(project)
                        write_log(project_id, "SUCCESS", f"连续生成完成！已生成到第{r_end2}集，请审稿")
                        return
                elif workflow.context.state == WorkflowState.COMPLETED:
                    project["status"] = "completed"
                    total_eps = project.get("current_episode") or project.get("total_episodes") or 80
                    review_start = review_anchor_start if review_anchor_start is not None else max(1, total_eps - batch_size + 1)
                    project["script_review_batch_start"] = review_start
                    project["script_review_batch_end"] = total_eps
                    try:
                        _sync_bible_from_version_history(project, (int(review_start), int(total_eps)))
                    except Exception:
                        pass
                    save_project(project)
                    write_log(project_id, "SUCCESS", f"🎉 连续生成全部完成！共生成 {project['current_episode']} 集")
                    return

            if auto_continue_to > 0:
                project["status"] = "checkpoint"
                review_start = review_anchor_start if review_anchor_start is not None else r_start
                project["script_review_batch_start"] = review_start
                project["script_review_batch_end"] = r_end
                save_project(project)
                write_log(project_id, "SUCCESS", f"连续生成完成！已生成到第{r_end}集，请审稿")
                return

            # 普通卡点：等人审稿
            project["status"] = "checkpoint"
            project["script_review_batch_start"] = r_start
            project["script_review_batch_end"] = r_end
            if "version_history" not in project:
                project["version_history"] = {}
            if "current_version" not in project:
                project["current_version"] = {}
            bible = workflow.context.bible
            if bible and bible.episodes:
                for ep_num in range(r_start, r_end + 1):
                    key = f"script_ep_{ep_num}"
                    if key not in project["version_history"]:
                        project["version_history"][key] = []
                    vers = project["version_history"][key]
                    if len(vers) > 1:
                        latest = dict(vers[-1])
                        latest["version_number"] = 1
                        latest["parent_version_index"] = -1
                        project["version_history"][key] = [latest]
                        project["current_version"][key] = 0
                        _update_process_storage_final(project, key, latest.get("content", ""), 0, latest.get("user_message", ""))
                    elif len(vers) == 1:
                        vers[0]["version_number"] = 1
                        vers[0]["parent_version_index"] = -1
                        project["current_version"][key] = 0
                        _update_process_storage_final(project, key, vers[0].get("content", ""), 0, vers[0].get("user_message", ""))
                    if len(project["version_history"][key]) == 0:
                        ep = next((e for e in bible.episodes if e.number == ep_num), None)
                        if ep and (ep.full_script or "").strip():
                            c = (ep.full_script or "").strip()
                            project["version_history"][key].append({
                                "content": c,
                                "timestamp": datetime.now().isoformat(),
                                "user_message": "",
                                "ai_response": "",
                                "version_number": 1,
                                "parent_version_index": -1,
                            })
                            project["current_version"][key] = 0
                            _update_process_storage_final(project, key, c, 0, "")
            save_project(project)
            write_log(project_id, "SUCCESS", f"本批第{r_start}-{r_end}集已生成，请审稿后点击「审核通过，继续下一批」")
            return
        
        if workflow.context.state == WorkflowState.COMPLETED:
            project["status"] = "completed"
            # 审稿范围优先沿用已记录范围（连续生成可能跨多批），否则回退到最后一批
            total_eps = project.get("current_episode") or project.get("total_episodes") or 80
            if project.get("script_review_batch_start") is not None and project.get("script_review_batch_end") is not None:
                r_start, _ = _get_script_review_range(project)
            else:
                batch_size = project.get("config", {}).get("batch_size", 5)
                r_start = max(1, total_eps - batch_size + 1)
            r_end = total_eps
            project["script_review_batch_start"] = r_start
            project["script_review_batch_end"] = r_end
            try:
                _sync_bible_from_version_history(project, (int(r_start), int(r_end)))
            except Exception:
                pass
            save_project(project)
            write_log(project_id, "SUCCESS", "=" * 60)
            write_log(project_id, "SUCCESS", f"🎉 全部完成！共生成 {project['current_episode']} 集")
            write_log(project_id, "SUCCESS", "=" * 60)
        else:
            project["status"] = "checkpoint"
            r_start = getattr(workflow.context, "review_batch_start", 0) or workflow.context.current_batch_start
            r_end = getattr(workflow.context, "review_batch_end", 0) or workflow.context.current_batch_end
            batch_size = project.get("config", {}).get("batch_size", 5)
            if r_start > r_end:
                r_start = max(1, r_end - batch_size + 1)
            project["script_review_batch_start"] = r_start
            project["script_review_batch_end"] = r_end
            # 人审前系统自动改的多版只留最新版；确保本批每集至少有一条首版
            if "version_history" not in project:
                project["version_history"] = {}
            if "current_version" not in project:
                project["current_version"] = {}
            bible = workflow.context.bible
            if bible and bible.episodes:
                for ep_num in range(r_start, r_end + 1):
                    key = f"script_ep_{ep_num}"
                    if key not in project["version_history"]:
                        project["version_history"][key] = []
                    vers = project["version_history"][key]
                    if len(vers) > 1:
                        # 人审前系统改了好几版：只留最新版
                        latest = dict(vers[-1])
                        latest["version_number"] = 1
                        latest["parent_version_index"] = -1
                        project["version_history"][key] = [latest]
                        project["current_version"][key] = 0
                        _update_process_storage_final(project, key, latest.get("content", ""), 0, latest.get("user_message", ""))
                    elif len(vers) == 1:
                        vers[0]["version_number"] = 1
                        vers[0]["parent_version_index"] = -1
                        project["current_version"][key] = 0
                        _update_process_storage_final(project, key, vers[0].get("content", ""), 0, vers[0].get("user_message", ""))
                    if len(project["version_history"][key]) == 0:
                        ep = next((e for e in bible.episodes if e.number == ep_num), None)
                        if ep and (ep.full_script or "").strip():
                            c = (ep.full_script or "").strip()
                            project["version_history"][key].append({
                                "content": c,
                                "timestamp": datetime.now().isoformat(),
                                "user_message": "",
                                "ai_response": "",
                                "version_number": 1,
                                "parent_version_index": -1,
                            })
                            project["current_version"][key] = 0
                            _update_process_storage_final(project, key, c, 0, "")
            save_project(project)
            write_log(project_id, "SUCCESS", f"本批第{r_start}-{r_end}集已生成，请审稿")

    except Exception as e:
        import traceback
        error_detail = traceback.format_exc()
        write_log(project_id, "ERROR", f"生成失败: {str(e)}")
        write_log(project_id, "ERROR", f"详细错误信息已记录")
        write_log(project_id, "ERROR", f"Traceback:\n{error_detail}")

        project = get_project(project_id)
        if project:
            project["status"] = "failed"
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
    finally:
        _thread_local_capture.current = None


def delete_project_by_id(project_id: str, username_key: str = "") -> bool:
    """删除项目及其输出目录"""
    path = get_project_path(project_id, username_key)
    if path.exists():
        path.unlink()
        import shutil
        if username_key:
            project_out = OUTPUT_BASE / username_key / project_id
        else:
            project_out = OUTPUT_BASE / project_id
        if project_out.exists():
            shutil.rmtree(project_out, ignore_errors=True)
        return True
    return False


def run_api():
    from fastapi import FastAPI, HTTPException, Request, Depends, Query
    from fastapi.responses import FileResponse, JSONResponse, StreamingResponse
    from fastapi.staticfiles import StaticFiles
    from fastapi.security import HTTPBearer, HTTPAuthorizationCredentials
    from pydantic import BaseModel
    import asyncio

    app = FastAPI(title="专业级标准化格式剧本生成 API")

    security = HTTPBearer(auto_error=False)

    # Token验证依赖：从 users.json 中查找用户
    async def verify_token(credentials: HTTPAuthorizationCredentials = Depends(security)):
        if not credentials:
            raise HTTPException(status_code=401, detail="未授权访问")
        token = (credentials.credentials or "").strip()
        user = find_user_by_token(token)
        if not user:
            raise HTTPException(status_code=401, detail="Token无效")
        return user

    # 静态资源与页面
    AXIS_WEB = Path(__file__).resolve().parent / "web"
    if AXIS_WEB.exists():
        app.mount("/assets", StaticFiles(directory=str(AXIS_WEB / "assets")), name="assets")

        @app.get("/")
        def root():
            return FileResponse(AXIS_WEB / "login.html")

        @app.get("/login")
        @app.get("/login.html")
        def login_page():
            return FileResponse(AXIS_WEB / "login.html")

        @app.get("/projects")
        @app.get("/projects.html")
        def projects_page():
            return FileResponse(AXIS_WEB / "projects.html")

        @app.get("/create")
        @app.get("/create.html")
        def create_page():
            return FileResponse(AXIS_WEB / "create.html")

        @app.get("/detail.html")
        def detail_page():
            return FileResponse(AXIS_WEB / "detail.html")


    class CreateBody(BaseModel):
        user_idea: str
        total_episodes: int = 80
        target_audience: str = "通用"
        batch_size: int = 5
        sample_strategy: str = "auto"
        provider: str = "wlai"
        checkpoint_after_ideation: bool = True

    class RegisterBody(BaseModel):
        username: str
        password: str

    class LoginBody(BaseModel):
        username: str = ""
        password: str = ""
        token: str = ""

    # 用户注册 / 登录 API
    @app.post("/api/auth/register")
    def api_register(body: RegisterBody):
        """注册新用户，自动生成 Token"""
        try:
            user = create_user(body.username, body.password)
        except ValueError as e:
            raise HTTPException(status_code=400, detail=str(e))
        return {
            "success": True,
            "username": user["username"],
            "token": user["token"],
            "message": "注册成功，已自动登录",
        }

    @app.post("/api/auth/login")
    def api_login(body: LoginBody):
        """用户名+密码登录，或使用 Token 登录"""
        import hashlib
        # Token 登录（兼容旧方式）
        if body.token and body.token.strip():
            user = find_user_by_token(body.token.strip())
            if not user:
                raise HTTPException(status_code=401, detail="Token 无效")
            return {
                "success": True,
                "username": user["username"],
                "token": user["token"],
                "message": "登录成功",
            }
        # 用户名+密码登录
        username = (body.username or "").strip()
        password = (body.password or "").strip()
        if not username or not password:
            raise HTTPException(status_code=400, detail="请输入用户名和密码")
        user = find_user_by_username(username)
        if not user:
            raise HTTPException(status_code=401, detail="用户名不存在，请先注册")
        stored_hash = user.get("password_hash", "")
        if not stored_hash:
            raise HTTPException(status_code=401, detail="该账号尚未设置密码，请联系管理员重置")
        input_hash = hashlib.sha256(password.encode("utf-8")).hexdigest()
        if input_hash != stored_hash:
            raise HTTPException(status_code=401, detail="密码错误")
        return {
            "success": True,
            "username": user["username"],
            "token": user["token"],
            "message": "登录成功",
        }

    @app.post("/api/auth/verify")
    async def api_verify(request: Request):
        """验证 Token 有效性"""
        auth_header = request.headers.get("Authorization", "")
        if not auth_header.startswith("Bearer "):
            return {"valid": False, "message": "无效的授权头"}
        token = auth_header.replace("Bearer ", "").strip()
        user = find_user_by_token(token)
        if not user:
            return {"valid": False, "message": "Token无效"}
        return {"valid": True, "message": "Token有效", "username": user["username"]}

    # ---- 用户自定义 LLM 配置档 API ----
    class AddModelBody(BaseModel):
        name: str
        api_key: str
        base_url: str = ""
        model: str = ""

    @app.get("/api/models", dependencies=[Depends(verify_token)])
    def api_list_models(user=Depends(verify_token)):
        from drama_agent.config import load_llm_profiles
        username_key = user.get("username_key")
        profiles = []
        for item in load_llm_profiles():
            if item.get("owner_username_key") != username_key:
                continue
            profiles.append({
                "id": f"profile__{item.get('id')}",
                "name": item.get("name") or "未命名配置",
                "model": item.get("model") or "",
                "base_url": item.get("base_url") or "",
                "custom": True,
            })
        return profiles

    @app.post("/api/models", dependencies=[Depends(verify_token)])
    def api_add_model(body: AddModelBody, user=Depends(verify_token)):
        name = (body.name or "").strip()
        api_key = (body.api_key or "").strip()
        base_url = (body.base_url or "").strip()
        model = (body.model or "").strip()
        if not name or not api_key or not base_url or not model:
            raise HTTPException(status_code=400, detail="配置名称、API Key、Base URL、模型名称不能为空")

        from drama_agent.config import load_llm_profiles, save_llm_profiles
        username_key = user.get("username_key")
        profiles = load_llm_profiles()
        if any(p.get("owner_username_key") == username_key and (p.get("name") or "").strip() == name for p in profiles):
            raise HTTPException(status_code=400, detail="该配置名称已存在")

        profile_id = str(uuid4())[:8]
        profiles.append({
            "id": profile_id,
            "name": name,
            "api_key": api_key,
            "base_url": base_url,
            "model": model,
            "owner_username": user.get("username"),
            "owner_username_key": username_key,
            "created_at": datetime.now().isoformat(),
        })
        save_llm_profiles(profiles)
        return {"success": True, "id": f"profile__{profile_id}", "message": f"配置 {name} 已保存"}

    @app.delete("/api/models/{model_id}", dependencies=[Depends(verify_token)])
    def api_delete_model(model_id: str, user=Depends(verify_token)):
        if not model_id.startswith("profile__"):
            raise HTTPException(status_code=400, detail="仅支持删除自定义配置")
        profile_id = model_id[len("profile__"):]
        from drama_agent.config import load_llm_profiles, save_llm_profiles
        profiles = load_llm_profiles()
        new_profiles = []
        removed = False
        for item in profiles:
            if str(item.get("id", "")) == profile_id and item.get("owner_username_key") == user.get("username_key"):
                removed = True
                continue
            new_profiles.append(item)
        if not removed:
            raise HTTPException(status_code=404, detail="未找到该配置")
        save_llm_profiles(new_profiles)
        return {"success": True, "message": "配置已删除"}

    @app.get("/api/projects")
    def api_list_projects(user=Depends(verify_token)):
        """项目列表"""
        username_key = user.get("username_key")
        return list_projects(owner_username_key=username_key)

    @app.get("/api/projects/{project_id}")
    def api_get_project(project_id: str, user=Depends(verify_token)):
        """获取项目详情"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权访问该项目")
        return project

    @app.get("/api/projects/{project_id}/logs")
    def api_get_project_logs(project_id: str, user=Depends(verify_token)):
        """获取项目日志"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权访问该项目")
        
        logs = get_project_logs(project_id)
        return {"project_id": project_id, "logs": logs}

    @app.get("/api/projects/{project_id}/bible")
    def api_get_project_bible(project_id: str, user=Depends(verify_token)):
        """获取项目的Bible数据"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权访问该项目")

        if project.get("script_review_batch_start") is not None and project.get("script_review_batch_end") is not None:
            try:
                _sync_bible_from_version_history(
                    project,
                    (
                        int(project.get("script_review_batch_start") or 1),
                        int(project.get("script_review_batch_end") or 1),
                    ),
                )
            except Exception:
                pass
        
        bible_path = _resolve_output_dir(project_id, user.get("username_key", "")) / "bible.json"
        if not bible_path.exists():
            return {}
        
        try:
            with open(bible_path, 'r', encoding='utf-8') as f:
                bible_data = json.load(f)
            return bible_data
        except Exception as e:
            raise HTTPException(status_code=500, detail=f"读取Bible失败: {str(e)}")
    
    @app.post("/api/projects/{project_id}/update")
    async def api_update_project(project_id: str, request: Request, user=Depends(verify_token)):
        """更新项目内容（支持修改灵感、标题、梗概等）"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权访问该项目")
        
        # 获取更新数据（使用 await 避免 asyncio.run 在已有事件循环中死锁）
        try:
            updates = await request.json()
        except Exception:
            raise HTTPException(status_code=400, detail="无效的请求数据")
        
        # 更新项目字段
        if 'inspiration' in updates:
            project['inspiration'] = updates['inspiration']
        
        # 更新Bible文件
        bible_path = _resolve_output_dir(project_id, user.get("username_key", "")) / "bible.json"
        if bible_path.exists():
            try:
                with open(bible_path, 'r', encoding='utf-8') as f:
                    bible_data = json.load(f)
                
                # 更新Bible内容
                if 'title' in updates:
                    bible_data['title'] = updates['title']
                    project['title'] = updates['title']
                
                if 'synopsis' in updates:
                    if not is_invalid_creative_response(updates['synopsis']):
                        bible_data['synopsis'] = updates['synopsis']
                    # 否则跳过，不把「编程助手拒绝」等无效内容写入 bible
                
                if 'genre' in updates:
                    # 将字符串转为列表
                    genre_str = updates['genre']
                    bible_data['genre'] = [g.strip() for g in genre_str.split('、') if g.strip()]
                
                if 'overall_outline' in updates:
                    bible_data['overall_outline'] = updates['overall_outline']

                if 'multi_outline' in updates:
                    bible_data['multi_episode_outline'] = updates['multi_outline']
                if 'multi_episode_outline' in updates:
                    bible_data['multi_episode_outline'] = updates['multi_episode_outline']

                if 'character' in updates:
                    char = updates['character']
                    raw_original = updates.get('character_original_name')
                    original_name = (raw_original or '').strip() or None  # 编辑前用的字典 key，用于替换而非新建
                    if isinstance(char, dict) and char.get('name'):
                        if 'characters' not in bible_data:
                            bible_data['characters'] = {}
                        # 合并规则：name 用新的；其他字段「新值非空用新，新值为空则继承旧值」
                        def _empty(v):
                            return v is None or v == '' or v == [] or v == {}
                        # 先取旧数据（不删），用旧数据做底合并新数据，再写入新 key，最后再删旧 key，避免先删导致合并缺字段
                        old_char = {}
                        if original_name:
                            if original_name in bible_data['characters']:
                                old_char = dict(bible_data['characters'][original_name])
                            else:
                                for k in list(bible_data['characters'].keys()):
                                    if (k or '').strip() == original_name:
                                        old_char = dict(bible_data['characters'][k])
                                        break
                        else:
                            new_name_key = (char.get('name') or '').strip()
                            old_char = dict(bible_data['characters'].get(new_name_key) or {})
                        merged = dict(old_char)
                        for k, v in char.items():
                            if not _empty(v):
                                merged[k] = v
                        merged['name'] = (char.get('name') or '').strip()
                        if merged['name']:
                            bible_data['characters'][merged['name']] = merged
                            # 新条目已写入，再删旧 key（改名时避免留两条）
                            if original_name and original_name != merged['name']:
                                if original_name in bible_data['characters']:
                                    bible_data['characters'].pop(original_name)
                                else:
                                    for k in list(bible_data['characters'].keys()):
                                        if (k or '').strip() == original_name:
                                            bible_data['characters'].pop(k)
                                            break
                                # 全库 relationships 里凡 target 为旧名的，改为新名
                                new_name = merged['name']
                                for other_key, other_char in list(bible_data['characters'].items()):
                                    rels = other_char.get('relationships')
                                    if not isinstance(rels, list):
                                        continue
                                    for r in rels:
                                        if isinstance(r, dict) and (r.get('target') or '').strip() == original_name:
                                            r['target'] = new_name

                if 'characters' in updates and isinstance(updates.get('characters'), dict):
                    bible_data['characters'] = updates['characters']

                if 'beat_sheet' in updates and isinstance(updates.get('beat_sheet'), dict) and isinstance(updates['beat_sheet'].get('episodes'), list):
                    bible_data['beat_sheet'] = updates['beat_sheet']
                
                # 保存更新后的Bible
                with open(bible_path, 'w', encoding='utf-8') as f:
                    json.dump(bible_data, f, ensure_ascii=False, indent=2)
                
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"更新Bible失败: {str(e)}")
        
        # 保存项目
        project['updated_at'] = datetime.now().isoformat()
        save_project(project)
        
        return {"success": True, "message": "更新成功"}
    
    class ChatRequest(BaseModel):
        message: str
        field: str  # 'inspiration', 'synopsis', 'outline' 等
        current_content: str
    
    @app.post("/api/projects/{project_id}/chat")
    async def api_chat_refine(project_id: str, body: ChatRequest, user=Depends(verify_token)):
        """AI对话优化内容（流式返回）"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权访问该项目")
        
        async def generate_response():
            """流式生成AI响应"""
            try:
                # 导入配置
                from drama_agent.config import Config, resolve_provider_to_config
                from openai import AsyncOpenAI
                
                config = Config()
                provider = project.get("config", {}).get("provider", "wlai")
                resolve_provider_to_config(config, provider)
                
                # 获取激活的配置
                active_config = config.llm.get_active_config()
                
                # 创建异步OpenAI客户端
                client = AsyncOpenAI(
                    api_key=active_config["api_key"],
                    base_url=active_config.get("base_url"),
                    timeout=120.0
                )
                
                # 构建提示词
                field_names = {
                    'inspiration': '创作灵感',
                    'synopsis': '故事梗概',
                    'outline': '总体大纲',
                    'overall_outline': '总体大纲',
                    'multi_outline': '多集大纲',
                    'character': '人物设定'
                }
                field_name = field_names.get(body.field, '内容')

                if body.field == 'characters':
                    system_msg = (
                        "你是专业的短剧编剧助手，只做人设修改。"
                        "你必须根据用户的修改建议，输出修改后的「全部角色」的 JSON 对象（一个 object，key 为角色名，value 为该角色的完整人设）。"
                        "若用户要求某人改名，必须同步更新所有人设的 relationships 里对该角色的 target 引用。不要输出任何解释、markdown 代码块或其它内容。"
                    )
                    prompt = f"""用户要修改全部人设，请根据用户建议一次性输出修改后的「全部角色」的 JSON。

当前全部人设（JSON object，key 为角色名）：
{body.current_content}

用户的修改建议（请尽量一次性说全所有修改点，如改名、改关系等）：
{body.message}

要求：
1. 输出必须是一个 JSON 对象，key 为人设名称，value 为完整人设（含 name, identity, archetype, personality, background, skills, relationships 等）
2. 若用户要求某人改名，请把该角色的 name 改为新名，并把所有人设的 relationships 里 target 为旧名的改为新名
3. 不要用 markdown 代码块包裹，不要加任何说明，直接输出 JSON
4. 保留所有未要求修改的角色和字段

修改后的全部人设 JSON："""
                elif body.field == 'beat_sheet':
                    system_msg = (
                        "你是专业的短剧编剧助手，只做分集大纲（本批）的修改。"
                        "你必须根据用户的修改建议，输出修改后的「本批分集」的 JSON 数组。"
                        "每个元素为对象，包含 episode（集数）、synopsis（剧情梗概）、ending_hook（结尾钩子）、hook_type（可选）。"
                        "不要输出任何解释、markdown 代码块或其它内容，直接输出 JSON 数组。"
                    )
                    prompt = f"""用户要修改本批分集大纲，请根据用户建议输出修改后的本批分集 JSON 数组。

当前本批分集（JSON 数组）：
{body.current_content}

用户的修改建议：
{body.message}

要求：
1. 输出必须是一个 JSON 数组，每项为 {{ "episode": 集数, "synopsis": "剧情梗概", "ending_hook": "结尾钩子", "hook_type": "可选" }}
2. 只输出本批的集数，不要增加或删除集数
3. 不要用 markdown 代码块包裹，不要加任何说明，直接输出 JSON 数组

修改后的本批分集 JSON 数组："""
                elif body.field == 'character':
                    system_msg = (
                        "你是专业的短剧编剧助手，只做人设修改。"
                        "你必须根据用户的修改建议，只输出修改后的「这一个角色」的完整 JSON 对象，不要输出任何解释、markdown 代码块或其它内容。"
                    )
                    prompt = f"""用户要修改以下单个人物设定，请根据用户建议只输出修改后的这一个角色的完整 JSON 对象。

当前该角色的 JSON：
{body.current_content}

用户的修改建议：
{body.message}

要求：
1. 输出必须是合法的单个 JSON 对象，与输入结构一致（包含 name, identity, archetype, personality, background, core_goal, memory_point, skills, relationships 等）
2. 必须保留 relationships 数组：若用户未要求改人物关系，原样保留；若用户要求改与某人的关系，只修改对应项
3. 不要用 markdown 代码块包裹，不要加任何说明
4. 直接输出 JSON 即可

修改后的该角色 JSON："""
                elif body.field.startswith("script_ep_"):
                    try:
                        ep_num = int(body.field.replace("script_ep_", "").strip())
                    except ValueError:
                        ep_num = 1
                    system_msg = (
                        "你是专业的短剧编剧助手。你只做「单集剧本」的修改。"
                        "你必须根据用户的修改意见，直接输出修改后的「完整一集剧本」正文，不要任何解释、不要 markdown 代码块。"
                        "格式要求：场景头如 1-1 客厅 日 内 或 ## 1-1 客厅；对话为「角色名：台词」；动作行以 △ 开头。"
                        "全中文，禁止出现英文字母。"
                    )
                    prompt = f"""用户对【第{ep_num}集】剧本的修改意见：
{body.message}

当前第{ep_num}集剧本：
{body.current_content}

请根据上述意见直接输出修改后的完整剧本正文。只输出剧本内容，不要加「修改后」「如下」等说明。"""
                else:
                    system_msg = (
                        "你是专业的短剧编剧助手，只做剧本/梗概/大纲等创意内容的修改。"
                        "你必须根据用户的修改建议，直接输出修改后的完整内容，不要输出任何解释、拒绝、编程或代码相关的内容。"
                        "若为剧本内容，必须全中文，禁止出现英文字母。"
                    )
                    prompt = f"""用户对当前的{field_name}提出了修改建议。

当前{field_name}内容：
{body.current_content}

用户的修改建议：
{body.message}

请根据用户的建议，重新生成优化后的{field_name}。要求：
1. 保持原有风格和核心内容
2. 融入用户提出的修改建议
3. 确保内容连贯、完整
4. 直接输出优化后的内容，不要添加任何解释或说明

优化后的{field_name}："""
                
                # 多轮对话：加载该 field 的 chat_history，模型会看到历史所有轮次
                history_raw = project.get("chat_history", {}).get(body.field, []) or []
                history_messages = []
                for h in history_raw:
                    if isinstance(h, dict):
                        if "role" in h and "content" in h:
                            history_messages.append({"role": h["role"], "content": h["content"] or ""})
                        elif "user_message" in h and "ai_response" in h:
                            history_messages.append({"role": "user", "content": h["user_message"] or ""})
                            history_messages.append({"role": "assistant", "content": h["ai_response"] or ""})
                # 剧本单集上下文长，只保留最近 N 轮，避免超长
                max_rounds = 10 if body.field.startswith("script_ep_") else 20
                if len(history_messages) > max_rounds * 2:
                    history_messages = history_messages[-(max_rounds * 2) :]
                messages = [{"role": "system", "content": system_msg}] + history_messages + [{"role": "user", "content": prompt}]
                
                # 云雾系 provider 将 system 合并到 user，避免代理注入的 system prompt 冲突
                from drama_agent.utils.llm_client import _merge_system_to_user, _YUNWU_PROVIDERS
                if config.llm.provider in _YUNWU_PROVIDERS:
                    messages = _merge_system_to_user(messages)
                
                accumulated_text = ""
                max_tok = 4000 if body.field == 'characters' else (6000 if body.field.startswith('script_ep_') else 2000)
                stream = await client.chat.completions.create(
                    model=active_config["model"],
                    messages=messages,
                    temperature=0.7,
                    max_tokens=max_tok,
                    stream=True
                )
                
                async for chunk in stream:
                    if not chunk.choices:
                        continue
                    content = getattr(chunk.choices[0].delta, "content", None)
                    if content:
                        accumulated_text += content
                        # 发送SSE格式的数据
                        yield f"data: {json.dumps({'chunk': content, 'accumulated': accumulated_text}, ensure_ascii=False)}\n\n"
                
                # 发送完成信号
                yield f"data: {json.dumps({'done': True, 'final_content': accumulated_text}, ensure_ascii=False)}\n\n"
                
                # 把本轮对话追加到 chat_history 和 process_storage.iteration_prompts
                try:
                    proj = get_project(project_id)
                    if proj and accumulated_text:
                        if "chat_history" not in proj:
                            proj["chat_history"] = {}
                        if body.field not in proj["chat_history"]:
                            proj["chat_history"][body.field] = []
                        proj["chat_history"][body.field].append({"role": "user", "content": prompt})
                        proj["chat_history"][body.field].append({"role": "assistant", "content": accumulated_text})
                        # 过程存储：追加迭代 prompt
                        if "process_storage" not in proj:
                            proj["process_storage"] = {}
                        if body.field not in proj["process_storage"]:
                            proj["process_storage"][body.field] = {"iteration_prompts": [], "final_selected": None}
                        proj["process_storage"][body.field]["iteration_prompts"].append({
                            "timestamp": datetime.now().isoformat(),
                            "user_prompt": prompt,
                            "ai_response": accumulated_text,
                        })
                        proj["updated_at"] = datetime.now().isoformat()
                        save_project(proj)
                except Exception as persist_err:
                    print(f"chat_history 追加失败（不影响本次回复）: {persist_err}", flush=True)
                
            except Exception as e:
                import traceback
                error_detail = traceback.format_exc()
                print(f"AI生成错误: {error_detail}", flush=True)
                error_msg = f"AI生成失败: {str(e)}"
                yield f"data: {json.dumps({'error': error_msg}, ensure_ascii=False)}\n\n"
        
        return StreamingResponse(
            generate_response(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no"
            }
        )
    
    class SaveVersionRequest(BaseModel):
        field: str  # 'inspiration', 'synopsis', 'overall_outline', 'multi_outline'
        content: str
        user_message: str  # 用户的修改建议
        ai_response: str  # AI的完整响应
        parent_version_index: Optional[int] = None  # 基于哪一版修改的；不传则视为基于上一版（线性）
    
    @app.post("/api/projects/{project_id}/save_version")
    def api_save_version(project_id: str, body: SaveVersionRequest, user=Depends(verify_token)):
        """保存内容版本到历史记录"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权访问该项目")
        
        # 确保版本历史结构存在
        if "version_history" not in project:
            project["version_history"] = {
                "inspiration": [],
                "synopsis": [],
                "overall_outline": [],
                "multi_outline": [],
                "characters": [],
                "beat_sheet": [],
            }
        
        if "chat_history" not in project:
            project["chat_history"] = {
                "inspiration": [],
                "synopsis": [],
                "overall_outline": [],
                "multi_outline": [],
                "characters": [],
                "beat_sheet": [],
            }
        
        # 拒绝明显无效的 AI 回复（如被路由成「编程助手」时的拒绝文），避免污染版本与 bible；剧本审稿 script_ep_ 不校验
        field = body.field
        if not field.startswith("script_ep_") and is_invalid_creative_response(body.ai_response):
            raise HTTPException(
                status_code=400,
                detail="AI 返回了无效内容（非剧本/梗概修改结果），请重试或更换模型后再保存。"
            )
        
        # 创建新版本
        versions_list = project.get("version_history", {}).get(field, [])
        current_version_count = len(versions_list)
        # 剧本审稿 script_ep_*：若上一条与当前 content+user_message 相同且时间在 30 秒内，视为重复写入（如双击/重试），不追加，避免版本历史出现两个一模一样的框
        if field.startswith("script_ep_") and versions_list:
            last = versions_list[-1]
            if (last.get("content") == body.content and last.get("user_message") == body.user_message):
                try:
                    last_ts = datetime.fromisoformat(last.get("timestamp", "")[:26].rstrip("Z"))
                    if (datetime.now() - last_ts).total_seconds() < 30:
                        new_index = len(versions_list) - 1
                        return {
                            "success": True,
                            "version_index": new_index,
                            "total_versions": len(versions_list),
                        }
                except Exception:
                    pass
        # 基于哪一版：显式传入则用，否则为「上一版」（线性）；首版为 -1
        parent_index = body.parent_version_index
        if parent_index is None:
            parent_index = current_version_count - 1 if current_version_count > 0 else -1
        
        version = {
            "content": body.content,
            "timestamp": datetime.now().isoformat(),
            "user_message": body.user_message,
            "ai_response": body.ai_response,
            "version_number": current_version_count + 1,  # 版本号从1开始
            "parent_version_index": parent_index,  # 用于分支：对话历史只显示从当前版回溯到根的这条链
        }
        
        # 确保该字段在版本历史中存在（兼容旧项目缺少 multi_outline 等新字段）
        if "version_history" not in project:
            project["version_history"] = {}
        if field not in project["version_history"]:
            project["version_history"][field] = []
        project["version_history"][field].append(version)
        
        new_index = len(project["version_history"][field]) - 1
        if "current_version" not in project:
            project["current_version"] = {"inspiration": -1, "synopsis": -1, "overall_outline": -1, "multi_outline": -1, "characters": -1, "beat_sheet": -1}
        if field not in project["current_version"]:
            project["current_version"][field] = -1
        project["current_version"][field] = new_index  # 新保存的版本即为当前版本

        # 过程存储：更新最后选定的版本
        if "process_storage" not in project:
            project["process_storage"] = {}
        if field not in project["process_storage"]:
            project["process_storage"][field] = {"iteration_prompts": [], "final_selected": None}
        project["process_storage"][field]["final_selected"] = {
            "content": body.content,
            "version_index": new_index,
            "timestamp": datetime.now().isoformat(),
            "user_message": body.user_message,
        }
        # 若 save 时带 user_message 且与上一轮 chat 不同，追加到 iteration_prompts
        if body.user_message or body.ai_response:
            prompts = project["process_storage"][field]["iteration_prompts"]
            last = prompts[-1] if prompts else None
            if not last or last.get("user_prompt") != body.user_message or last.get("ai_response") != body.ai_response:
                project["process_storage"][field]["iteration_prompts"].append({
                    "timestamp": datetime.now().isoformat(),
                    "user_prompt": body.user_message,
                    "ai_response": body.ai_response,
                })

        # 保存项目
        project['updated_at'] = datetime.now().isoformat()
        save_project(project)

        return {
            "success": True,
            "version_index": new_index,
            "total_versions": len(project["version_history"][field])
        }
    
    @app.get("/api/projects/{project_id}/versions/{field}")
    def api_get_versions(project_id: str, field: str, user=Depends(verify_token)):
        """获取指定字段的版本历史"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权访问该项目")
        
        # 确保版本历史结构存在
        if "version_history" not in project:
            project["version_history"] = {
                "inspiration": [],
                "synopsis": [],
                "overall_outline": [],
                "multi_outline": [],
                "characters": [],
                "beat_sheet": [],
            }
        
        if field not in project["version_history"]:
            project["version_history"][field] = []
        versions = project["version_history"][field]
        current_version = project.get("current_version", {}).get(field, -1)
        if current_version == -1 and versions:
            current_version = len(versions) - 1
        
        return JSONResponse(
            content={
                "versions": versions,
                "current_version": current_version,
                "chat_history": project.get("chat_history", {}).get(field, []),
                "process_storage": project.get("process_storage", {}).get(field, {"iteration_prompts": [], "final_selected": None}),
            },
            headers={"Cache-Control": "no-cache, no-store, must-revalidate", "Pragma": "no-cache"}
        )
    
    class SelectVersionRequest(BaseModel):
        field: str
        version_index: int
    
    @app.post("/api/projects/{project_id}/select_version")
    def api_select_version(project_id: str, body: SelectVersionRequest, user=Depends(verify_token)):
        """选择某个版本作为当前版本"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权访问该项目")
        
        # 确保版本历史结构存在
        if "version_history" not in project or body.field not in project["version_history"]:
            raise HTTPException(status_code=400, detail="字段不存在")
        
        versions = project["version_history"][body.field]
        if body.version_index < 0 or body.version_index >= len(versions):
            raise HTTPException(status_code=400, detail="版本索引无效")
        
        # 更新当前版本索引
        if "current_version" not in project:
            project["current_version"] = {
                "inspiration": -1,
                "synopsis": -1,
                "overall_outline": -1,
                "multi_outline": -1,
                "characters": -1,
                "beat_sheet": -1,
            }
        if body.field not in project["current_version"]:
            project["current_version"][body.field] = -1
        project["current_version"][body.field] = body.version_index
        # 过程存储：更新最后选定的版本
        sel_ver = versions[body.version_index]
        if "process_storage" not in project:
            project["process_storage"] = {}
        if body.field not in project["process_storage"]:
            project["process_storage"][body.field] = {"iteration_prompts": [], "final_selected": None}
        project["process_storage"][body.field]["final_selected"] = {
            "content": sel_ver.get("content", ""),
            "version_index": body.version_index,
            "timestamp": datetime.now().isoformat(),
            "user_message": sel_ver.get("user_message", ""),
        }
        project['updated_at'] = datetime.now().isoformat()
        save_project(project)

        return {
            "success": True,
            "selected_content": versions[body.version_index]["content"]
        }

    @app.post("/api/projects")
    def api_create_project(body: CreateBody, user=Depends(verify_token)):
        """创建项目"""
        if not (body.user_idea and body.user_idea.strip()):
            raise HTTPException(status_code=400, detail="请输入生成提示词")
        project = create_project(
            user_idea=body.user_idea,
            total_episodes=min(100, max(5, body.total_episodes)),
            target_audience=body.target_audience,
            batch_size=min(10, max(1, body.batch_size)),
            sample_strategy=body.sample_strategy,
            provider=body.provider,
            checkpoint_after_ideation=body.checkpoint_after_ideation,
            username_key=user.get("username_key", ""),
        )
        # 绑定项目归属用户，确保不同用户项目隔离
        project["owner_username"] = user.get("username")
        project["owner_username_key"] = user.get("username_key")
        save_project(project)
        return {"project_id": project["id"], "title": project["title"]}

    @app.delete("/api/projects/{project_id}")
    def api_delete_project(project_id: str, user=Depends(verify_token)):
        """删除项目"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权删除该项目")
        if delete_project_by_id(project_id, username_key=user.get("username_key", "")):
            return {"success": True}
        raise HTTPException(status_code=404, detail="项目不存在")

    @app.post("/api/projects/{project_id}/start")
    def api_start_project(project_id: str, user=Depends(verify_token)):
        """启动项目 - 开始第一步：生成创作灵感"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权操作该项目")
        
        if project["status"] == "running":
            raise HTTPException(status_code=400, detail="项目正在生成中")
        
        # 检查是否启用了分步模式
        if project.get("config", {}).get("checkpoint_after_ideation", True):
            # 分步模式：只执行第一步
            project["status"] = "running"
            project["ideation_stage"] = "not_started"
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
            
            # 启动第一步：生成灵感
            thread = threading.Thread(
                target=run_ideation_step,
                args=(project_id, "inspiration"),
                daemon=True
            )
            thread.start()
            
            return {"success": True, "message": "开始生成创作灵感"}
        else:
            # 快速模式：直接运行完整流程
            project["status"] = "running"
            project["workflow_state"] = "running"
            project["current_episode"] = 0
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
            
            thread = threading.Thread(target=run_generation_task, args=(project_id,), daemon=True)
            thread.start()
            
            return {"success": True, "message": "已启动生成"}
    
    @app.get("/api/projects/{project_id}/logs/stream")
    async def api_stream_logs(project_id: str, tail: bool = False, user=Depends(verify_token)):
        """流式获取项目日志。tail=1 时只推送连接之后的新日志（用于「继续」下一步时不重放历史）。"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权访问该项目")
        
        log_file = _resolve_output_dir(project_id, user.get("username_key", "")) / "generation.log"
        
        async def log_generator():
            """生成日志流：完整行单独推送；未完成行以 append 增量推送，前端无换行追加，既流式又不拆行。"""
            last_position = 0
            line_tail = ""  # 当前未完成行（尚未读到 \\n）
            sent_partial = ""  # 当前未完成行里已通过 append 发出去的前缀，用于算增量
            if tail and log_file.exists():
                last_position = log_file.stat().st_size
            retry_count = 0
            max_retry = 120
            
            while retry_count < max_retry:
                try:
                    if log_file.exists():
                        with open(log_file, 'r', encoding='utf-8') as f:
                            f.seek(last_position)
                            new_content = f.read()
                            last_position = f.tell()
                            
                            if new_content or line_tail:
                                retry_count = 0
                                full = line_tail + new_content
                                parts = full.split('\n')
                                if full and not full.endswith('\n'):
                                    line_tail = parts.pop()
                                else:
                                    line_tail = ""
                                # 先发完整行，再发未完成行的 append，保证顺序且能区分「新行」
                                last_yielded_complete = False
                                for line in parts:
                                    line = line.strip()
                                    if line:
                                        yield f"data: {json.dumps({'log': line}, ensure_ascii=False)}\n\n"
                                        last_yielded_complete = True
                                    sent_partial = ""
                                # 未完成行：刚发过完整行时，前缀 \\n，这样同一 div 内会换行显示
                                if line_tail and line_tail != sent_partial:
                                    delta = line_tail[len(sent_partial):] if line_tail.startswith(sent_partial) else line_tail
                                    if delta:
                                        if last_yielded_complete:
                                            delta = "\n" + delta
                                        yield f"data: {json.dumps({'log': delta, 'append': True}, ensure_ascii=False)}\n\n"
                                    sent_partial = line_tail
                            else:
                                # 没有新内容，检查项目状态
                                current_project = get_project(project_id)
                                if current_project and current_project.get('status') in ['completed', 'failed', 'checkpoint']:
                                    yield f"data: {json.dumps({'status': current_project['status'], 'done': True}, ensure_ascii=False)}\n\n"
                                    break
                    
                    await asyncio.sleep(1)
                    retry_count += 1
                    
                except Exception as e:
                    yield f"data: {json.dumps({'error': str(e)}, ensure_ascii=False)}\n\n"
                    break
            
            if line_tail.strip():
                yield f"data: {json.dumps({'log': line_tail}, ensure_ascii=False)}\n\n"
            yield f"data: {json.dumps({'done': True}, ensure_ascii=False)}\n\n"
        
        return StreamingResponse(
            log_generator(),
            media_type="text/event-stream",
            headers={
                "Cache-Control": "no-cache",
                "Connection": "keep-alive",
                "X-Accel-Buffering": "no"
            }
        )
    
    @app.post("/api/projects/{project_id}/continue")
    def api_continue_project(project_id: str, user=Depends(verify_token)):
        """继续下一步（创意阶段各步骤、剧本撰写、或剧本审稿通过继续下一批）"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权操作该项目")
        # 允许 checkpoint；撰写失败后也可重试（status=failed 时一律允许继续）
        current_status = project.get("status")
        if current_status not in ("checkpoint", "failed"):
            raise HTTPException(
                status_code=400,
                detail=f"项目不在卡点状态（当前 status={current_status!r}，仅允许 checkpoint 或 failed）"
            )
        
        # 剧本审稿卡点：审核通过，继续撰写下一批（先读出批次范围再 pop，传入线程跑记忆官）
        if project.get("script_review_batch_end") is not None:
            r_start = project.get("script_review_batch_start")
            r_end = project.get("script_review_batch_end")
            project["status"] = "running"
            project.pop("script_review_batch_start", None)
            project.pop("script_review_batch_end", None)
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
            thread = threading.Thread(
                target=run_scripting_task,
                args=(project_id,),
                kwargs={"script_review_batch_start": r_start, "script_review_batch_end": r_end},
                daemon=True,
            )
            thread.start()
            return {"success": True, "message": "已审核通过，开始撰写下一批"}
        
        ideation_stage = project.get("ideation_stage", "not_started")
        # 撰写失败重试时若阶段未记录，视为创意已完成、直接开始撰写
        if project["status"] == "failed" and ideation_stage not in (
            "inspiration", "synopsis", "characters", "overall_outline", "multi_outline", "beat_sheet", "completed"
        ):
            ideation_stage = "completed"
        
        # 根据当前阶段决定下一步（beat_sheet 每批审核后仍下一步为 beat_sheet，直到全部完成）
        stage_map = {
            "inspiration": ("synopsis", "生成故事梗概"),
            "synopsis": ("characters", "生成人物角色"),
            "characters": ("overall_outline", "生成总体大纲"),
            "overall_outline": ("multi_outline", "生成多集大纲"),
            "multi_outline": ("beat_sheet", "生成分集大纲（第1批）"),
            "beat_sheet": ("beat_sheet", "生成下一批分集大纲"),
            "completed": (None, "开始撰写剧本")
        }
        
        if ideation_stage not in stage_map:
            raise HTTPException(status_code=400, detail=f"未知的阶段: {ideation_stage}")
        
        next_step, message = stage_map[ideation_stage]
        
        project["status"] = "running"
        project["updated_at"] = datetime.now().isoformat()
        save_project(project)
        
        if next_step:
            # 继续创意阶段的下一步
            thread = threading.Thread(
                target=run_ideation_step,
                args=(project_id, next_step),
                daemon=True
            )
            thread.start()
            return {"success": True, "message": f"开始{message}"}
        else:
            # 创意阶段完成，开始剧本撰写（第一批）
            thread = threading.Thread(target=run_scripting_task, args=(project_id,), daemon=True)
            thread.start()
            return {"success": True, "message": "开始撰写剧本"}

    # ---- 连续生成到指定集数 ----
    class ContinueToBody(BaseModel):
        target_episode: int

    @app.post("/api/projects/{project_id}/continue_to")
    def api_continue_to(project_id: str, body: ContinueToBody, user=Depends(verify_token)):
        """跳过审稿，连续生成到指定集数"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权操作该项目")
        current_status = project.get("status")
        if current_status not in ("checkpoint", "failed"):
            raise HTTPException(status_code=400, detail="项目不在卡点状态")
        if project.get("script_review_batch_end") is None:
            raise HTTPException(status_code=400, detail="当前不在剧本撰写卡点")
        target = body.target_episode
        total = project.get("total_episodes") or 80
        if target < 1 or target > total:
            raise HTTPException(status_code=400, detail=f"目标集数需在 1-{total} 之间")
        r_start = project.get("script_review_batch_start")
        r_end = project.get("script_review_batch_end")
        project["status"] = "running"
        project.pop("script_review_batch_start", None)
        project.pop("script_review_batch_end", None)
        project["updated_at"] = datetime.now().isoformat()
        save_project(project)
        thread = threading.Thread(
            target=run_scripting_task,
            args=(project_id,),
            kwargs={"script_review_batch_start": r_start, "script_review_batch_end": r_end, "auto_continue_to": target},
            daemon=True,
        )
        thread.start()
        return {"success": True, "message": f"已开始连续生成到第{target}集"}

    # ---- 一次性生成全部分集大纲 ----
    @app.post("/api/projects/{project_id}/beat_sheet_all")
    def api_beat_sheet_all(project_id: str, user=Depends(verify_token)):
        """一次性生成剩余所有分集大纲（不中途卡点）"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权操作该项目")
        current_status = project.get("status")
        if current_status not in ("checkpoint", "failed"):
            raise HTTPException(status_code=400, detail="项目不在卡点状态")
        ideation_stage = project.get("ideation_stage", "not_started")
        if ideation_stage not in ("multi_outline", "beat_sheet"):
            raise HTTPException(status_code=400, detail=f"当前阶段({ideation_stage})不支持此操作，需在多集大纲或分集大纲审核阶段")
        project["status"] = "running"
        project["updated_at"] = datetime.now().isoformat()
        save_project(project)
        thread = threading.Thread(
            target=run_ideation_step,
            args=(project_id, "beat_sheet"),
            kwargs={"generate_all": True},
            daemon=True,
        )
        thread.start()
        total = project.get("total_episodes", 80)
        next_start = project.get("beat_sheet_next_start", 1)
        return {"success": True, "message": f"已开始一次性生成分集大纲（第{next_start}-{total}集）"}

    class ScriptReviewFeedbackBody(BaseModel):
        feedback: str
        episode_number: Optional[int] = None  # 若传则只重写该集（按集多轮对话）

    @app.post("/api/projects/{project_id}/script_review/feedback")
    def api_script_review_feedback(project_id: str, body: ScriptReviewFeedbackBody, user=Depends(verify_token)):
        """剧本审稿：提交修改意见。传 episode_number 则只重写该集；不传则重写本批全部。卡点或已完成最后一批均可。"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权操作该项目")
        in_review = project.get("script_review_batch_end") is not None
        if project["status"] == "completed":
            total = project.get("total_episodes") or project.get("current_episode") or 80
            in_review = in_review or (project.get("current_episode") or 0) >= total
        if not in_review or project["status"] not in ("checkpoint", "completed"):
            raise HTTPException(status_code=400, detail="当前不在剧本审稿卡点或本批无需审稿")
        feedback = (body.feedback or "").strip()
        if not feedback:
            raise HTTPException(status_code=400, detail="请输入修改意见")
        batch_start = project.get("script_review_batch_start", 1)
        batch_end = project.get("script_review_batch_end", 1)
        if project["status"] == "completed":
            batch_start, batch_end = _get_script_review_range(project)
        elif batch_start > batch_end:
            batch_start, batch_end = batch_end, batch_start
        bible_path = None
        if project.get("bible_path") and Path(project["bible_path"]).exists():
            bible_path = Path(project["bible_path"])
        if bible_path is None and (_resolve_output_dir(project_id, user.get("username_key", "")) / "bible.json").exists():
            bible_path = _resolve_output_dir(project_id, user.get("username_key", "")) / "bible.json"
        if bible_path is None or not bible_path.exists():
            raise HTTPException(status_code=404, detail="未找到剧本数据")
        output_dir = bible_path.parent
        from drama_agent.workflow.drama_workflow import DramaWorkflow
        from drama_agent.config import Config, resolve_provider_to_config, set_config
        from drama_agent.utils.llm_client import reset_llm_client
        config = Config()
        config.output_dir = str(output_dir)
        config.bible_path = str(bible_path)
        config.drama.episodes_per_batch = project.get("config", {}).get("batch_size", 5)
        provider = project.get("config", {}).get("provider", "wlai")
        resolve_provider_to_config(config, provider)
        set_config(config)
        reset_llm_client()
        workflow = DramaWorkflow()
        workflow.config = config
        workflow.resume(str(bible_path))
        ep_num = body.episode_number
        if ep_num is not None:
            # 按集：只重写该集，只更新该集版本历史
            if not (batch_start <= ep_num <= batch_end):
                raise HTTPException(status_code=400, detail="episode_number 不在本批范围内")
            ok = workflow.apply_human_script_feedback_for_episode(ep_num, feedback)
            if not ok:
                raise HTTPException(status_code=400, detail="未找到该集或重写失败")
            if workflow.context.bible:
                workflow.context.bible.save(str(bible_path))
                if "version_history" not in project:
                    project["version_history"] = {}
                key = f"script_ep_{ep_num}"
                if key not in project["version_history"]:
                    project["version_history"][key] = []
                vers = project["version_history"][key]
                parent = len(vers) - 1 if vers else -1
                ep = next((e for e in workflow.context.bible.episodes if e.number == ep_num), None)
                new_content = (ep.full_script or "") if ep else ""
                vers.append({
                    "content": new_content,
                    "timestamp": datetime.now().isoformat(),
                    "user_message": feedback,
                    "ai_response": "",
                    "version_number": len(vers) + 1,
                    "parent_version_index": parent,
                })
                _update_process_storage_final(project, key, new_content, len(vers) - 1, feedback)
            project["updated_at"] = datetime.now().isoformat()
            save_project(project)
            return {"success": True, "message": f"已根据修改意见重写第{ep_num}集，请刷新查看"}
        # 按批：重写本批全部
        workflow.apply_human_script_feedback(batch_start, batch_end, feedback)
        if workflow.context.bible:
            workflow.context.bible.save(str(bible_path))
            if "version_history" not in project:
                project["version_history"] = {}
            bible = workflow.context.bible
            for ep in bible.episodes:
                if not (batch_start <= ep.number <= batch_end):
                    continue
                key = f"script_ep_{ep.number}"
                if key not in project["version_history"]:
                    project["version_history"][key] = []
                vers = project["version_history"][key]
                parent = len(vers) - 1 if vers else -1
                new_content = ep.full_script or ""
                vers.append({
                    "content": new_content,
                    "timestamp": datetime.now().isoformat(),
                    "user_message": feedback,
                    "ai_response": "",
                    "version_number": len(vers) + 1,
                    "parent_version_index": parent,
                })
                _update_process_storage_final(project, key, new_content, len(vers) - 1, feedback)
        project["updated_at"] = datetime.now().isoformat()
        save_project(project)
        return {"success": True, "message": "已根据修改意见重写本批剧本，请刷新查看"}

    class ScriptEpisodeSaveBody(BaseModel):
        episode_number: int
        full_script: str

    @app.put("/api/projects/{project_id}/script_episode")
    def api_save_script_episode(project_id: str, body: ScriptEpisodeSaveBody, user=Depends(verify_token)):
        """剧本审稿：保存单集剧本（手动编辑），写入该集 version_history。卡点或已完成均可保存最后一批。"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权操作该项目")
        # 允许：卡点审稿 或 已完成时保存最后一批（76-80）
        in_review = project.get("script_review_batch_end") is not None
        if project["status"] == "completed":
            total = project.get("total_episodes") or project.get("current_episode") or 80
            in_review = in_review or (project.get("current_episode") or 0) >= total
        if not in_review or (project["status"] not in ("checkpoint", "completed")):
            raise HTTPException(status_code=400, detail="当前不在剧本审稿卡点")
        batch_start = project.get("script_review_batch_start", 1)
        batch_end = project.get("script_review_batch_end", 1)
        if project["status"] == "completed":
            batch_start, batch_end = _get_script_review_range(project)
        elif batch_start > batch_end:
            batch_start, batch_end = batch_end, batch_start
        ep_num = body.episode_number
        if not (batch_start <= ep_num <= batch_end):
            raise HTTPException(status_code=400, detail="episode_number 不在本批范围内")
        output_dir = _resolve_output_dir(project_id, user.get("username_key", ""))
        bible_path = output_dir / "bible.json"
        if not bible_path.exists():
            raise HTTPException(status_code=404, detail="未找到剧本数据")
        from drama_agent.models import Bible
        bible = Bible.load(str(bible_path))
        episode = next((ep for ep in bible.episodes if ep.number == ep_num), None)
        if not episode:
            raise HTTPException(status_code=404, detail="未找到该集")
        script_text = (body.full_script or "").strip()
        episode.full_script = script_text
        bible.save(str(bible_path))
        if "version_history" not in project:
            project["version_history"] = {}
        key = f"script_ep_{ep_num}"
        if key not in project["version_history"]:
            project["version_history"][key] = []
        vers = project["version_history"][key]
        # 幂等：若上一条已是「手动编辑」且内容相同、30 秒内，不重复追加
        if vers:
            last = vers[-1]
            if (last.get("user_message") == "手动编辑" and last.get("content") == script_text):
                try:
                    last_ts = datetime.fromisoformat(last.get("timestamp", "")[:26].rstrip("Z"))
                    if (datetime.now() - last_ts).total_seconds() < 30:
                        project["updated_at"] = datetime.now().isoformat()
                        save_project(project)
                        return {"success": True, "message": f"已保存第{ep_num}集剧本"}
                except Exception:
                    pass
        parent = len(vers) - 1 if vers else -1
        vers.append({
            "content": script_text,
            "timestamp": datetime.now().isoformat(),
            "user_message": "手动编辑",
            "ai_response": "",
            "version_number": len(vers) + 1,
            "parent_version_index": parent,
        })
        _update_process_storage_final(project, key, script_text, len(vers) - 1, "手动编辑")
        project["updated_at"] = datetime.now().isoformat()
        save_project(project)
        return {"success": True, "message": f"已保存第{ep_num}集剧本"}

    class ApplyScriptContentBody(BaseModel):
        episode_number: int
        content: str
        user_message: str = ""

    @app.post("/api/projects/{project_id}/script_episode/apply_content")
    def api_apply_script_content(project_id: str, body: ApplyScriptContentBody, user=Depends(verify_token)):
        """剧本审稿：将流式重写得到的内容写入该集并加入版本历史（与 save_version 一致，不做卡点校验）"""
        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        owner_key = project.get("owner_username_key")
        if owner_key and owner_key != user.get("username_key"):
            raise HTTPException(status_code=403, detail="无权操作该项目")
        ep_num = body.episode_number
        output_dir = _resolve_output_dir(project_id, user.get("username_key", ""))
        bible_path = output_dir / "bible.json"
        if not bible_path.exists():
            raise HTTPException(status_code=404, detail="未找到剧本数据")
        from drama_agent.models import Bible
        bible = Bible.load(str(bible_path))
        episode = next((ep for ep in bible.episodes if ep.number == ep_num), None)
        if not episode:
            raise HTTPException(status_code=404, detail="未找到该集")
        script_text = (body.content or "").strip()
        episode.full_script = script_text
        bible.save(str(bible_path))
        return {"success": True, "message": f"已更新第{ep_num}集剧本"}

    @app.get("/api/projects/{project_id}/export")
    def api_export_project(project_id: str, token: str = Depends(verify_token)):
        """导出项目为 Word 文档：标题、梗概、大纲、人设、分集剧本。四号楷体、1.5倍行距、主标题加粗、对白首行缩进2字符。"""
        import io
        import urllib.parse
        from fastapi.responses import StreamingResponse
        from docx import Document
        from docx.shared import Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
        from docx.oxml.ns import qn

        def set_run_font(run, font_name="楷体", font_size_pt=14, bold=False):
            try:
                run.font.name = font_name
                run.font.size = Pt(font_size_pt)
                run.font.bold = bold
                rPr = getattr(run._element, "rPr", None)
                if rPr is not None:
                    rFonts = getattr(rPr, "rFonts", None)
                    if rFonts is not None:
                        rFonts.set(qn("w:eastAsia"), font_name)
            except Exception:
                pass

        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        if (project.get("current_episode") or 0) < 1:
            raise HTTPException(status_code=400, detail="尚无已生成的剧集，无法导出")

        # 优先用项目里保存的 bible_path（绝对路径），再试 ROOT/output/pid
        bible_file = None
        try:
            if project.get("bible_path"):
                bp = Path(project["bible_path"])
                if bp.exists():
                    bible_file = bp
        except Exception:
            pass
        if bible_file is None and (_resolve_output_dir(project_id, user.get("username_key", "")) / "bible.json").exists():
            bible_file = _resolve_output_dir(project_id, user.get("username_key", "")) / "bible.json"
        if bible_file is None:
            raise HTTPException(status_code=400, detail="未找到剧本数据(bible.json)，无法导出")

        try:
            with open(bible_file, "r", encoding="utf-8") as f:
                raw = json.load(f)
            bible_data = raw if isinstance(raw, dict) else {}
        except Exception as e:
            raise HTTPException(status_code=500, detail="读取剧本文件失败: " + str(e).replace("\n", " "))

        try:
            doc = Document()
            font_pt = 14  # 四号
            line_rule = WD_LINE_SPACING.ONE_POINT_FIVE

            # 默认正文样式（不依赖 rPr 存在，避免部分环境报错）
            try:
                style = doc.styles["Normal"]
                style.font.name = "楷体"
                style.font.size = Pt(font_pt)
                style.paragraph_format.line_spacing_rule = line_rule
                if getattr(style._element, "rPr", None) is not None and getattr(style._element.rPr, "rFonts", None) is not None:
                    style._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
            except Exception:
                pass

            def add_para(text, first_line_indent_pt=None, bold=False):
                p = doc.add_paragraph()
                p.paragraph_format.line_spacing_rule = line_rule
                if first_line_indent_pt is not None:
                    p.paragraph_format.first_line_indent = Pt(first_line_indent_pt)
                r = p.add_run(text or "")
                set_run_font(r, font_size_pt=font_pt, bold=bold)
                return p

            def add_heading_text(text, level=1, bold=True):
                h = doc.add_heading(text, level=level)
                for run in h.runs:
                    set_run_font(run, font_size_pt=font_pt, bold=bold)
                return h

            # 1. 标题（加粗，用段落兼容各版本 docx）
            title_text = project.get("title") or bible_data.get("title") or "未命名项目"
            title = doc.add_paragraph(title_text)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                title.style = doc.styles["Title"]
            except Exception:
                pass
            for run in title.runs:
                set_run_font(run, font_size_pt=font_pt, bold=True)
            doc.add_paragraph()

            # 2. 梗概（主标题加粗）
            add_heading_text("梗概", level=1, bold=True)
            synopsis = bible_data.get("synopsis") or project.get("user_idea") or "无"
            add_para(synopsis)
            doc.add_paragraph()

            # 3. 大纲（主标题加粗）
            add_heading_text("大纲", level=1, bold=True)
            outline = bible_data.get("overall_outline") or bible_data.get("multi_episode_outline") or "无"
            add_para(outline)
            doc.add_paragraph()

            # 4. 人设（主标题加粗，每人名加粗；兼容 char 为 dict 或 str）
            add_heading_text("人设", level=1, bold=True)
            characters = bible_data.get("characters")
            if isinstance(characters, dict):
                for name, char in characters.items():
                    add_heading_text(str(name), level=2, bold=True)
                    if isinstance(char, str):
                        if char.strip():
                            add_para(char.strip())
                    elif isinstance(char, dict):
                        profile_parts = [
                            char.get("identity"),
                            char.get("personality"),
                            char.get("background"),
                            char.get("profile"),
                        ]
                        profile_text = " ".join(str(x).strip() for x in profile_parts if x)
                        if profile_text:
                            add_para(profile_text)
            doc.add_paragraph()

            # 5. 分集剧本（第X集加粗，对白首行缩进2字符）
            episodes = bible_data.get("episodes") or []
            if not isinstance(episodes, list):
                episodes = []
            if not episodes:
                add_heading_text("分集剧本", level=1, bold=True)
                add_para("暂无分集剧本。")
            else:
                # 按集数排序，仅处理元素为 dict 的项
                episodes = sorted([e for e in episodes if isinstance(e, dict)], key=lambda e: e.get("number", 0))
                for ep in episodes:
                    doc.add_page_break()
                    ep_num = ep.get("number", 0)
                    add_heading_text(f"第{ep_num}集", level=1, bold=True)
                    script_content = ep.get("full_script") or ep.get("script") or ""
                    if script_content and isinstance(script_content, str):
                        for line in script_content.split("\n"):
                            line = line.strip()
                            if not line:
                                continue
                            # 1. ## 去掉，后面内容加粗；2. △ 开头一律不缩进；3. 场景下「人物：」一律不缩进
                            if line.startswith("##"):
                                display_text = line.lstrip("#").strip()
                                if display_text:
                                    add_para(display_text, first_line_indent_pt=None, bold=True)
                                continue
                            if line.startswith("△"):
                                add_para(line, first_line_indent_pt=None, bold=False)
                                continue
                            if line.startswith("人物："):
                                add_para(line, first_line_indent_pt=None, bold=False)
                                continue
                            is_dialogue = "：" in line or (":" in line)
                            indent_pt = Pt(28) if is_dialogue else None
                            p = doc.add_paragraph()
                            p.paragraph_format.line_spacing_rule = line_rule
                            if indent_pt is not None:
                                p.paragraph_format.first_line_indent = indent_pt
                            r = p.add_run(line)
                            set_run_font(r, font_size_pt=font_pt, bold=False)

            doc_buffer = io.BytesIO()
            doc.save(doc_buffer)
            doc_buffer.seek(0)
        except Exception as e:
            logging.exception("导出 Word 失败")
            err_msg = str(e).replace("\n", " ").encode("ascii", errors="replace").decode("ascii")
            raise HTTPException(status_code=500, detail="Export failed: " + err_msg)

        filename = f"{project.get('title', '短剧剧本')}_{project_id}.docx"
        return StreamingResponse(
            doc_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{urllib.parse.quote(filename)}"},
        )

    @app.get("/api/projects/{project_id}/export_ideation")
    def api_export_ideation(project_id: str, token: str = Depends(verify_token)):
        """导出创意策划完整版为 Word（灵感、梗概、人设详情、总体大纲、多集大纲、分集大纲）"""
        from fastapi.responses import StreamingResponse
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_LINE_SPACING
        from docx.oxml.ns import qn
        import io
        import urllib.parse

        project = get_project(project_id)
        if not project:
            raise HTTPException(status_code=404, detail="项目不存在")
        ideation = project.get("ideation_stage", "")
        if ideation not in ("beat_sheet", "completed"):
            raise HTTPException(status_code=400, detail="创意策划未完成，无法导出")

        output_dir = _resolve_output_dir(project_id, user.get("username_key", ""))
        bible_path = output_dir / "bible.json"
        if bible_path.exists():
            with open(bible_path, "r", encoding="utf-8") as f:
                bible_data = json.load(f)
        else:
            # bible.json 可能尚未生成或路径不一致：仅用 project 内数据组 doc，避免 404
            bible_data = {
                "title": project.get("title") or "未命名项目",
                "synopsis": "",
                "characters": {},
                "overall_outline": "",
                "multi_episode_outline": "",
                "multi_outline": "",
                "beat_sheet": {"episodes": []},
                "total_episodes": project.get("total_episodes", 80),
            }

        doc = Document()
        style = doc.styles["Normal"]
        style.font.name = "楷体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
        style.font.size = Pt(14)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

        def add_para(text, bold=False):
            p = doc.add_paragraph(text or "")
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            for run in p.runs:
                run.font.name = "楷体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
                run.font.size = Pt(14)
                if bold:
                    run.bold = True
            return p

        def add_heading1(title):
            h = doc.add_heading(title, level=1)
            for run in h.runs:
                run.font.name = "楷体"
                run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
            return h

        title = project.get("title") or bible_data.get("title") or "未命名项目"
        add_heading1("创意策划书")
        add_para(f"剧本标题：{title}", bold=True)
        add_para(f"规划集数：{project.get('total_episodes') or bible_data.get('total_episodes') or 80} 集")
        doc.add_paragraph()

        # 1. 创作灵感
        add_heading1("一、创作灵感")
        inspiration = (project.get("inspiration") or "").strip()
        add_para(inspiration if inspiration else "（无）")
        doc.add_paragraph()

        # 2. 故事梗概
        add_heading1("二、故事梗概")
        synopsis = (bible_data.get("synopsis") or "").strip()
        add_para(synopsis if synopsis else "（无）")
        doc.add_paragraph()

        # 3. 人设（完整）
        add_heading1("三、人物设定")
        chars_raw = bible_data.get("characters") or {}
        chars = list(chars_raw.values()) if isinstance(chars_raw, dict) else list(chars_raw)
        if not chars:
            add_para("（无）")
        else:
            for c in chars:
                name = c.get("name") or "未命名"
                doc.add_heading(name, level=2)
                for run in doc.paragraphs[-1].runs:
                    run.font.name = "楷体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
                add_para(f"身份：{c.get('identity') or ''}")
                add_para(f"性格：{c.get('personality') or ''}")
                add_para(f"背景：{c.get('background') or ''}")
                if c.get("core_goal"):
                    add_para(f"核心目标：{c['core_goal']}")
                if c.get("memory_point"):
                    add_para(f"记忆点：{c['memory_point']}")
                rels = c.get("relationships") or []
                if rels:
                    rel_str = "；".join([
                        (f"{r.get('target', '')} — {r.get('relation_type') or r.get('relation') or ''}") for r in rels
                    ])
                    add_para(f"关系：{rel_str}")
                doc.add_paragraph()

        # 4. 总体大纲
        add_heading1("四、总体大纲")
        overall = (bible_data.get("overall_outline") or "").strip()
        add_para(overall if overall else "（无）")
        doc.add_paragraph()

        # 5. 多集大纲
        add_heading1("五、多集大纲")
        multi = (bible_data.get("multi_episode_outline") or bible_data.get("multi_outline") or "").strip()
        add_para(multi if multi else "（无）")
        doc.add_paragraph()

        # 6. 分集大纲
        add_heading1("六、分集大纲")
        beat_sheet = bible_data.get("beat_sheet") or {}
        episodes = beat_sheet.get("episodes") or []
        if not episodes:
            add_para("（无）")
        else:
            for b in sorted(episodes, key=lambda x: x.get("episode") or 0):
                ep = b.get("episode", "?")
                doc.add_heading(f"第{ep}集", level=2)
                for run in doc.paragraphs[-1].runs:
                    run.font.name = "楷体"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
                add_para(f"梗概：{b.get('synopsis') or ''}")
                add_para(f"结尾钩子：{b.get('ending_hook') or ''}")
                if b.get("hook_type"):
                    add_para(f"钩子类型：{b['hook_type']}")
                doc.add_paragraph()

        doc_buffer = io.BytesIO()
        doc.save(doc_buffer)
        doc_buffer.seek(0)
        safe_name = urllib.parse.quote(f"{title}_创意策划.docx")
        return StreamingResponse(
            doc_buffer,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f"attachment; filename*=UTF-8''{safe_name}"},
        )

    import uvicorn
    port = int(os.environ.get("AXIS_API_PORT", "8502"))
    uvicorn.run(app, host="0.0.0.0", port=port)


if __name__ == "__main__":
    run_api()
